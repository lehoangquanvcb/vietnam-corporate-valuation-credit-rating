import sys as _sys
from pathlib import Path as _Path
_PROJECT_ROOT = _Path(__file__).resolve().parents[1]
if str(_PROJECT_ROOT) not in _sys.path:
    _sys.path.insert(0, str(_PROJECT_ROOT))

import numpy as np
from pathlib import Path
from io import BytesIO
import math, re, pandas as pd, matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scripts.universal_data import get_company,get_snapshot,entity_history,industry_metric_history,industry_label,num
from scripts.multisector_valuation import valuation
from scripts.three_methodology_rating import rate_company
from scripts.rating_committee_engine import committee_pack
from scripts.fair_value_range import fair_value_range
from scripts.valuation_triangulation import triangulate
from scripts.rating_evidence_engine import rating_evidence
from scripts.intelligent_analyst import analyze as intelligent_analyze
from scripts.sector_kpi_engine import sector_kpi_table
from scripts.sector_templates import get_template
from scripts.methodology_kpi_engine import methodology_kpi_table, metric_list, LABELS as METH_LABELS, PCT as METH_PCT, MULT as METH_MULT
from scripts.public_intelligence import load_public_intelligence, scope_for_entity

ROOT=Path(__file__).resolve().parents[1]

VI_METRIC={
'TotalAssets':'Tổng tài sản','Revenue':'Doanh thu','NPAT':'Lợi nhuận sau thuế','ROE':'ROE','ROA':'ROA',
'NPL':'Tỷ lệ nợ xấu','CAR':'CAR','CASA':'CASA','LDR':'LDR','NIM':'NIM','PB':'P/B','PE':'P/E',
'DebtEquity':'Nợ/VCSH','CurrentRatio':'Hệ số thanh toán hiện hành','AvailableCapitalRatio':'Tỷ lệ vốn khả dụng',
'AssetEquity':'Tổng tài sản/VCSH','CreditCostProxy':'Chi phí dự phòng/Dư nợ','FundingGapAssets':'Funding gap/TTS',
'CashAssets':'Tiền/TTS','WorkingCapitalAssets':'VLĐ ròng/TTS','NetDebtEquity':'Nợ ròng/VCSH','NetDebtEBITDA':'Nợ ròng/EBITDA',
'EquityAssetsCorp':'VCSH/TTS','AssetTurnover':'Vòng quay tài sản','FOCFMargin':'FOCF/Doanh thu','CashDebt':'Tiền/Nợ vay'
}
PCT={'ROE','ROA','NPL','CAR','CASA','NIM','AvailableCapitalRatio','CreditCostProxy','FundingGapAssets','CashAssets','WorkingCapitalAssets','EquityAssetsCorp','FOCFMargin','CashDebt'}
MULT={'PB','PE','DebtEquity','CurrentRatio','LDR','AssetEquity','NetDebtEquity','NetDebtEBITDA','AssetTurnover'}

ENTITY_TYPE_VI={'BANK':'Ngân hàng','SECURITIES':'Công ty Chứng khoán','CORPORATE':'Doanh nghiệp Phi tài chính'}

def entity_type_vi(v):
    key=str(v or '').strip().upper()
    return ENTITY_TYPE_VI.get(key, str(v or 'N/A'))

def company_display_name(meta,ticker):
    name=str(meta.get('DisplayName') or meta.get('LegalName') or meta.get('CompanyName') or ticker).strip()
    # Always show the ticker in parentheses exactly once.
    if re.search(r'\('+re.escape(str(ticker))+r'\)\s*$',name,re.I):
        return name
    return f"{name} ({str(ticker).upper()})"

def _public_intel_paragraphs(entity_type, only=None):
    scope=scope_for_entity(entity_type)
    items=load_public_intelligence(scope)
    out=[]
    for x in items:
        if only=='macro' and str(x.get('Scope')).upper()!='MACRO': continue
        if only=='industry' and str(x.get('Scope')).upper()=='MACRO': continue
        title=str(x.get('Title') or '').strip(); nar=str(x.get('Narrative') or '').strip()
        if nar: out.append((title,nar,str(x.get('Source') or ''),str(x.get('URL') or ''),str(x.get('AsOf') or '')))
    return out

def compact_abs(v):
    x=num(v)
    if x is None:return 'N/A'
    ax=abs(x)
    if ax>=1e12:return (f"{x/1e12:,.1f} nghìn tỷ").replace(',','X').replace('.',',').replace('X','.')
    if ax>=1e9:return (f"{x/1e9:,.1f} tỷ").replace(',','X').replace('.',',').replace('X','.')
    if ax>=1e6:return (f"{x/1e6:,.1f} triệu").replace(',','X').replace('.',',').replace('X','.')
    return vi(x,1)

def vi(x,d=1):
    v=num(x)
    if v is None:return 'N/A'
    return f'{v:,.{d}f}'.replace(',','X').replace('.',',').replace('X','.')
def pct(x): return 'N/A' if num(x) is None else vi(num(x)*100,1)+'%'
def mult(x): return 'N/A' if num(x) is None else vi(x,2)+'x'
def price(x): return 'N/A' if num(x) is None else vi(num(x)*1000,0)+' đồng/cp'
def metric_fmt(metric,x):
    if metric in PCT:return pct(x)
    if metric in MULT:return mult(x)
    return compact_abs(x)

def _font():
    try:return font_manager.findfont('Lato',fallback_to_default=False)
    except:return 'DejaVu Sans'

def _period_date(v):
    s=str(v)
    for q,m,d in [('Q1','03','31'),('Q2','06','30'),('Q3','09','30'),('Q4','12','31')]:
        s=s.replace(q,f'-{m}-{d}')
    return pd.to_datetime(s,errors='coerce')

def chart_metric(ticker,metric,title=None,percent=None):
    h=entity_history(ticker); p=industry_metric_history(ticker,metric)
    fig,ax=plt.subplots(figsize=(8.4,3.15))
    plt.rcParams.update({'font.family':'Lato','font.size':10})
    z=pd.DataFrame()
    if len(h):
        z=h[h.Metric.astype(str).eq(metric)].copy()
        z['Date']=z.Period.map(_period_date); z['Value']=pd.to_numeric(z.Value,errors='coerce')
        z=z.dropna(subset=['Date','Value']).sort_values('Date').drop_duplicates('Date',keep='last')
        if len(z): ax.plot(z.Date,z.Value,marker='o',linewidth=1.8,label=str(ticker).upper())
    if len(p):
        pp=p.copy()
        if 'PeriodDate' in pp: pp=pp.sort_values('PeriodDate')
        ax.plot(pp.PeriodDate,pp.IndustryMean,linestyle='--',linewidth=1.6,label=industry_label(ticker))
    ax.set_title(title or f"{VI_METRIC.get(metric,metric)} - doanh nghiệp và trung bình ngành",fontsize=11)
    ax.grid(alpha=.2); ax.legend(fontsize=8,loc='best')
    if percent is None: percent=metric in PCT
    if percent: ax.yaxis.set_major_formatter(lambda v,pos:(f'{v*100:.1f}%').replace('.',','))
    fig.autofmt_xdate(rotation=0); fig.tight_layout()
    bio=BytesIO(); fig.savefig(bio,dpi=180,bbox_inches='tight'); plt.close(fig); bio.seek(0); return bio

def _set_cell_shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def _set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def _set_cell_margins(cell,top=60,start=70,bottom=60,end=70):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def _style_doc(doc, report_type=None):
    sec=doc.sections[0]; sec.page_height=Mm(297); sec.page_width=Mm(210)
    sec.top_margin=Mm(16); sec.bottom_margin=Mm(15); sec.left_margin=Mm(18); sec.right_margin=Mm(16)
    for name in ['Normal','Title','Subtitle','Heading 1','Heading 2','Heading 3']:
        st=doc.styles[name]; st.font.name='Lato'
        st._element.rPr.rFonts.set(qn('w:ascii'),'Lato'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Lato'); st._element.rPr.rFonts.set(qn('w:eastAsia'),'Lato')
    n=doc.styles['Normal']; n.font.size=Pt(11); n.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    n.paragraph_format.line_spacing=1.15; n.paragraph_format.space_after=Pt(4)
    doc.styles['Title'].font.size=Pt(20); doc.styles['Title'].font.bold=True
    doc.styles['Heading 1'].font.size=Pt(14); doc.styles['Heading 1'].font.bold=True
    doc.styles['Heading 1'].paragraph_format.space_before=Pt(8); doc.styles['Heading 1'].paragraph_format.space_after=Pt(5)
    doc.styles['Heading 2'].font.size=Pt(12); doc.styles['Heading 2'].font.bold=True
    doc.styles['Heading 2'].paragraph_format.space_before=Pt(6); doc.styles['Heading 2'].paragraph_format.space_after=Pt(4)
    # Header is report-specific and functions as a clear simulation/disclaimer notice.
    header=sec.header.paragraphs[0]
    if report_type == 'rating':
        header.text='Báo cáo mô phỏng quá trình Xếp hạng tín nhiệm (tài liệu này không thể thay thế Báo cáo Xếp hạng tín nhiệm)'
    elif report_type == 'analysis':
        header.text='Báo cáo mô phỏng quá trình Phân tích, định giá cổ phiếu (tài liệu này không nhằm mục đích khuyến nghị đầu tư cổ phiếu)'
    else:
        header.text=''
    header.alignment=WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_before=Pt(0); header.paragraph_format.space_after=Pt(0)
    for r in header.runs:
        r.font.name='Lato'; r.font.size=Pt(9); r.bold=False
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run('Tài liệu chỉ có tính chất tham khảo, tác giả không chịu trách nhiệm về độ chính xác của dữ liệu đầu vào cũng như kết quả đầu ra')
    r.font.name='Lato'; r.font.size=Pt(9)

def _add_title(doc,ticker,meta,report_type):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('BÁO CÁO XẾP HẠNG TÍN NHIỆM' if report_type=='rating' else 'BÁO CÁO PHÂN TÍCH GIÁ CỔ PHIẾU, ĐỊNH GIÁ & M&A')
    r.bold=True;r.font.name='Lato';r.font.size=Pt(20)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(company_display_name(meta,ticker));r.bold=True;r.font.name='Lato';r.font.size=Pt(16)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Ngành: {meta.get('Sector')} | Nhóm so sánh: {meta.get('PeerGroup')}").font.size=Pt(10)
    doc.add_paragraph()
    box=doc.add_table(rows=3,cols=2);box.alignment=WD_TABLE_ALIGNMENT.CENTER;box.style='Table Grid'
    vals=[('Loại hình',entity_type_vi(meta.get('EntityType'))),('Phương pháp',meta.get('Methodology')),('Nguồn benchmark',industry_label(ticker))]
    for i,(a,b) in enumerate(vals):
        box.cell(i,0).text=str(a);box.cell(i,1).text=str(b)
        _set_cell_shading(box.cell(i,0),'E8EEF7')
    doc.add_page_break()

def _add_toc(doc,sections):
    doc.add_heading('MỤC LỤC NỘI DUNG',0)
    for i,(h,_) in enumerate(sections,1):
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(2)
        p.add_run(f'{i}. {h}')
    doc.add_page_break()

def _add_kpi_table(doc,ticker,limit=10):
    k,_,_=sector_kpi_table(ticker)
    if not len(k):return
    t=doc.add_table(rows=1,cols=5);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=['Chỉ tiêu','Doanh nghiệp','TB ngành','Trung vị ngành','Số DN']
    for j,x in enumerate(hdr):
        t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),'E8EEF7')
    _set_repeat_table_header(t.rows[0])
    for _,r in k.head(limit).iterrows():
        c=t.add_row().cells;m=r['Metric']
        vals=[r['Chỉ tiêu'],metric_fmt(m,r['Doanh nghiệp']),metric_fmt(m,r['Trung bình ngành']),metric_fmt(m,r['Trung vị ngành']),str(int(r['Số DN có dữ liệu']))]
        for j,x in enumerate(vals):c[j].text=str(x)
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;_set_cell_margins(cell)
            for p in cell.paragraphs:
                for run in p.runs:run.font.name='Lato';run.font.size=Pt(10)

def _analyst_paragraphs(ticker):
    a=intelligent_analyze(ticker)
    pars=[]
    pars.append(a.get('Conclusion',''))
    if a.get('Strengths'): pars.append('Điểm mạnh tương đối: '+' '.join(a['Strengths'][:3]))
    if a.get('Risks'): pars.append('Rủi ro/điểm yếu tương đối: '+' '.join(a['Risks'][:3]))
    if a.get('Interpretations'): pars.append('Nhận định chéo: '+' '.join(a['Interpretations'][:2]))
    return pars

def _stock_sections(ticker,meta,s,val):
    fv=fair_value_range(ticker); _,tmpl=get_template(meta.get('EntityType'),meta.get('Sector'))
    return [
    ('TÓM TẮT ĐIỀU HÀNH',[
        f"{meta.get('CompanyName')} được phân tích theo mẫu chuyên ngành {tmpl['label']}, với benchmark {industry_label(ticker)}.",
        *_analyst_paragraphs(ticker),
        f"Vùng giá mô hình hiện tại: Bear {price(fv.get('Bear'))}; Base {price(fv.get('Base'))}; Bull {price(fv.get('Bull'))}; kịch bản chiến lược/M&A {price(fv.get('StrategicMA'))}. Kịch bản M&A phải được hiệu chỉnh theo từng giao dịch cụ thể."
    ]),
    ('HỒ SƠ DOANH NGHIỆP & NGÀNH',[
        f"Doanh nghiệp thuộc ngành {meta.get('Sector')}. Các trọng tâm chuyên ngành gồm: {', '.join(tmpl['focus'])}.",
        f"Phương pháp định giá ưu tiên theo template: {', '.join(tmpl['valuation'])}. Benchmark được tính từ doanh nghiệp cùng ngành có dữ liệu ở từng chỉ tiêu/kỳ."
    ]),
    ('PHÂN TÍCH KẾT QUẢ KINH DOANH & TĂNG TRƯỞNG',[
        "Đánh giá tăng trưởng cần tách tăng trưởng quy mô khỏi chất lượng tăng trưởng. Doanh thu/tổng tài sản được đặt cạnh trung bình ngành để nhận diện doanh nghiệp đang mở rộng nhanh hơn hay chậm hơn peer.",
        "Kết luận chính thức cần kiểm tra nguồn tăng trưởng, tính lặp lại, cơ cấu doanh thu/tài sản và các khoản mục bất thường."
    ]),
    ('KHẢ NĂNG SINH LỢI & HIỆU QUẢ VỐN',[
        f"ROE hiện tại {pct(s.get('ROE'))}; ROA {pct(s.get('ROA'))}. Các tỷ suất này được so sánh đồng thời với trung bình và trung vị ngành.",
        "ROE cao không tự động đồng nghĩa định giá cao hơn: cần kiểm tra đòn bẩy, chất lượng lợi nhuận, vòng quay vốn và mức độ bền vững."
    ]),
    ('BẢNG CÂN ĐỐI, ĐÒN BẨY & THANH KHOẢN',[
        "Phân tích tập trung vào khả năng hấp thụ tổn thất, cơ cấu nguồn vốn, đòn bẩy và thanh khoản. Chỉ tiêu đặc thù thay đổi theo ngân hàng, CTCK và doanh nghiệp phi tài chính.",
        "Các chỉ tiêu được đọc theo xu hướng thời gian và tương quan peer; không kết luận chỉ từ một kỳ đơn lẻ."
    ]),
    ('SO SÁNH NHÓM TƯƠNG ĐỒNG',[
        f"Nhóm so sánh hiện tại: {industry_label(ticker)}. Mỗi KPI hiển thị số doanh nghiệp thực sự có dữ liệu để tránh tạo cảm giác chính xác giả khi coverage thấp.",
        "Trung vị ngành được dùng như kiểm tra chéo với trung bình ngành trong trường hợp outlier lớn."
    ]),
    ('ĐỊNH GIÁ CƠ SỞ & CHẾ ĐỘ ĐỊNH GIÁ',[
        f"Giá trị tham chiếu của engine: {price(val.get('FairValue'))}; phương pháp chính: {val.get('PrimaryMethod','N/A')}.",
        "Định giá được cross-check với P/E, P/B, EV/EBITDA hoặc phương pháp nội tại phù hợp loại hình. Kết quả phải được đọc cùng ROE, tăng trưởng và rủi ro."
    ]),
    ('BEAR - BASE - BULL',[
        f"Bear {price(fv.get('Bear'))}; Base {price(fv.get('Base'))}; Bull {price(fv.get('Bull'))}.",
        "Bear/Base/Bull là khung độ nhạy, không phải ba mức giá chắc chắn. Analyst có thể override haircut/premium dựa trên stress test và triển vọng ngành."
    ]),
    ('M&A, QUYỀN KIỂM SOÁT & GIÁ TRỊ CHIẾN LƯỢC',[
        f"Kịch bản Strategic/M&A hiện tại {price(fv.get('StrategicMA'))}, dựa trên control premium và synergy scenario của từng case.",
        "Không áp dụng vùng giá hoặc control premium của STB cho các doanh nghiệp khác. Với giao dịch thực tế cần phân tích quy mô lô, quyền kiểm soát, scarcity, synergy, nguồn vốn và điều kiện giao dịch."
    ]),
    ('STRESS TEST, CATALYST & RỦI RO ĐẦU TƯ',[
        "Stress test phải liên kết biến số ngành với doanh thu/lợi nhuận, vốn, thanh khoản và multiple định giá. Catalyst cần có khả năng mở khóa giá trị và mốc thời gian kiểm chứng được.",
        *_analyst_paragraphs(ticker)[1:]
    ]),
    ('KẾT LUẬN PHÂN TÍCH',[
        "Kết luận cuối cùng tổng hợp fundamental, vị trí tương đối so với ngành, định giá cơ sở, vùng Bear/Base/Bull và giá trị chiến lược/M&A.",
        "Kết quả của platform là công cụ hỗ trợ phân tích; quyết định đầu tư cần cập nhật dữ liệu thị trường, thông tin doanh nghiệp và giả định tại thời điểm sử dụng."
    ]),
    ('PHỤ LỤC DỮ LIỆU, PEER & GIẢ ĐỊNH',[
        "Phụ lục ghi nhận KPI, coverage peer, nguồn dữ liệu và các giả định/override. Đây là audit trail để tái lập kết quả."
    ])
    ]

def _rating_sections(ticker,meta,s,rr):
    rc=committee_pack(ticker); method=rr.get('MethodologyName')
    return [
    ('TÓM TẮT XẾP HẠNG',[
        f"Phương pháp được tự động lựa chọn: {method}. Kết quả mô phỏng: Anchor {rr.get('Anchor','N/A')}; SACP/SCA {rr.get('SACP',rr.get('SCA','N/A'))}; ICR {rr.get('ICR','N/A')}.",
        "Kết quả tự động là đầu vào hỗ trợ chuyên viên/Hội đồng, không thay thế phê duyệt XHTN chính thức."
    ]),
    ('PHẠM VI & PHƯƠNG PHÁP LUẬN',[
        rr.get('Audit',''),
        "Platform giữ riêng ba methodology: Ngân hàng, Công ty chứng khoán và Doanh nghiệp phi tài chính; không dùng một scorecard chung cho ba loại hình."
    ]),
    ('RỦI RO VĨ MÔ & NGÀNH',[
        f"Doanh nghiệp hoạt động trong ngành {meta.get('Sector')}. Đánh giá ngành được đặt trong bối cảnh chu kỳ, cạnh tranh, rào cản và khả năng truyền dẫn rủi ro vào hồ sơ tín dụng.",
        "Các yếu tố định tính cần được chuyên viên cập nhật từ hồ sơ doanh nghiệp và nguồn thẩm định."
    ]),
    ('HỒ SƠ KINH DOANH',[
        "Đánh giá vị thế cạnh tranh, quy mô, đa dạng hóa, franchise, mô hình kinh doanh và chất lượng tăng trưởng. Peer benchmark được dùng như bằng chứng định lượng hỗ trợ.",
        "Các KCF chưa có dữ liệu cấu trúc được giữ dưới dạng judgment/override thay vì tự sinh số."
    ]),
    ('HỒ SƠ TÀI CHÍNH',[
        f"ROE {pct(s.get('ROE'))}; ROA {pct(s.get('ROA'))}. Phân tích tài chính được điều chỉnh theo loại hình doanh nghiệp và đặt cạnh peer.",
        "Trọng tâm gồm vốn/đòn bẩy, khả năng sinh lợi, rủi ro tài sản, nguồn vốn, thanh khoản và dòng tiền."
    ]),
    ('ANCHOR, NOTCH/MODIFIER & SACP/SCA',[
        f"Anchor {rr.get('Anchor','N/A')}; SACP/SCA {rr.get('SACP',rr.get('SCA','N/A'))}. Waterfall dưới đây cho phép truy vết từng bước từ điểm khởi đầu đến năng lực tín nhiệm độc lập.",
        "Mọi Analyst Override cần lưu lý do, nguồn dữ liệu và người phê duyệt."
    ]),
    ('HỖ TRỢ BÊN NGOÀI & ICR',[
        f"Hỗ trợ bên ngoài hiện tại {rr.get('ExternalSupportNotches',0)} bậc; ICR {rr.get('ICR','N/A')}.",
        "Chỉ ghi nhận hỗ trợ khi có cơ sở về năng lực và động cơ hỗ trợ theo methodology áp dụng."
    ]),
    ('ĐỘ NHẠY XẾP HẠNG & STRESS TEST',[
        "Độ nhạy cần chỉ ra điều kiện cụ thể có thể nâng/hạ hạng: thay đổi vốn, chất lượng tài sản, đòn bẩy, thanh khoản, lợi nhuận, vị thế kinh doanh hoặc hỗ trợ.",
        "Stress test phải liên kết kịch bản bất lợi với các chỉ tiêu dẫn tới thay đổi notch/modifier hoặc rating cap."
    ]),
    ('EARLY WARNING & GIÁM SÁT SAU XẾP HẠNG',[
        "Thiết lập chỉ báo cảnh báo sớm theo methodology và ngành. Khi chỉ báo vượt ngưỡng, chuyên viên cần đánh giá lại giả định và khả năng thay đổi rating.",
        "Không dùng một ngưỡng chung cho mọi ngành."
    ]),
    ('KẾT LUẬN TRÌNH HỘI ĐỒNG XHTN',[
        f"Kết quả mô phỏng hiện tại: {rr.get('ICR','N/A')}. Hội đồng cần xem xét đầy đủ waterfall, peer, dữ liệu nguồn, override và sensitivity trước khi phê duyệt.",
        "Báo cáo chính thức phải phân biệt rõ kết quả máy tính, nhận định chuyên viên và quyết định Hội đồng."
    ]),
    ('PHỤ LỤC KPI & PEER',["Bảng KPI và benchmark dùng để kiểm tra tính nhất quán của các luận điểm định lượng."]),
    ('PHỤ LỤC WATERFALL & AUDIT TRAIL',["Waterfall ghi lại từng cấu phần, kết quả, điều chỉnh và luận cứ để phục vụ tái kiểm tra."])
    ]

def _add_pars(doc,pars):
    for txt in pars:
        if not txt:continue
        p=doc.add_paragraph(str(txt));p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

def _add_waterfall(doc,ticker,components=None,compact=False):
    rc=committee_pack(ticker);rows=rc.get('Waterfall',[])
    if components:
        wanted={str(x).strip().lower() for x in components}
        rows=[r for r in rows if str(r.get('Cấu phần','')).strip().lower() in wanted]
    if not rows:return
    cols=4 if compact else 5
    t=doc.add_table(rows=1,cols=cols);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=(['Cấu phần','Kết quả','Điều chỉnh','Luận cứ chính'] if compact else ['Bước','Cấu phần','Kết quả','Điều chỉnh','Luận cứ'])
    for j,x in enumerate(hdr):t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),'E8EEF7')
    if not compact:_set_repeat_table_header(t.rows[0])
    for r in rows:
        c=t.add_row().cells
        if compact:
            vals=[r.get('Cấu phần',''),r.get('Kết quả',''),r.get('Điều chỉnh',''),str(r.get('Luận cứ','')).split(';')[0][:90]]
        else:
            vals=[r.get(k,'') for k in ['Bước','Cấu phần','Kết quả','Điều chỉnh','Luận cứ']]
        for j,x in enumerate(vals):c[j].text=str(x)
    for row in t.rows:
        trPr=row._tr.get_or_add_trPr();cant=OxmlElement('w:cantSplit');trPr.append(cant)
        for cell in row.cells:
            _set_cell_margins(cell,top=30,start=45,bottom=30,end=45)
            for pp in cell.paragraphs:
                pp.paragraph_format.space_after=Pt(0);pp.paragraph_format.line_spacing=1.0
                for run in pp.runs:run.font.name='Lato';run.font.size=Pt(8 if compact else 8.5)

def _add_fv_table(doc,ticker):
    f=fair_value_range(ticker)
    t=doc.add_table(rows=2,cols=4);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=['Bear','Base','Bull','Chiến lược/M&A']; vals=[price(f.get('Bear')),price(f.get('Base')),price(f.get('Bull')),price(f.get('StrategicMA'))]
    for j,x in enumerate(hdr):t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),'E8EEF7')
    for j,x in enumerate(vals):t.cell(1,j).text=x
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:run.font.name='Lato';run.font.size=Pt(10)


def professional_page_plan(report_type):
    if report_type=='rating':
        return [
        ('TÓM TẮT XẾP HẠNG','Kết quả xếp hạng, triển vọng và luận điểm tín dụng cốt lõi.'),
        ('PHẠM VI XẾP HẠNG & DỮ LIỆU','Phạm vi tổ chức phát hành, BCTC hợp nhất, kỳ dữ liệu và nguồn benchmark.'),
        ('PHƯƠNG PHÁP LUẬN ÁP DỤNG','Tự động route đúng methodology Ngân hàng / CTCK / Doanh nghiệp phi tài chính.'),
        ('RỦI RO VĨ MÔ','Các yếu tố kinh tế vĩ mô có khả năng truyền dẫn vào hồ sơ tín dụng.'),
        ('RỦI RO NGÀNH','Chu kỳ, cấu trúc cạnh tranh, rào cản và mức độ biến động của ngành.'),
        ('HỒ SƠ KINH DOANH','Vị thế cạnh tranh, franchise, quy mô, đa dạng hóa và chất lượng tăng trưởng.'),
        ('QUY MÔ & TĂNG TRƯỞNG','Tăng trưởng tài sản/doanh thu và mức độ bền vững so với peer.'),
        ('KHẢ NĂNG SINH LỢI','ROE, ROA, biên lợi nhuận và chất lượng nguồn thu.'),
        ('CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN','NPL hoặc các rủi ro tài sản đặc thù theo loại hình.'),
        ('VỐN & ĐÒN BẨY','Mức độ đầy đủ vốn, đòn bẩy và bộ đệm hấp thụ tổn thất.'),
        ('NGUỒN VỐN','Cơ cấu nguồn vốn, mức độ ổn định, tập trung và chi phí.'),
        ('THANH KHOẢN','Khả năng đáp ứng nghĩa vụ và nguồn thanh khoản dự phòng.'),
        ('DÒNG TIỀN & KHẢ NĂNG TRẢ NỢ','Khả năng chuyển lợi nhuận thành tiền và đáp ứng nghĩa vụ nợ.'),
        ('QUẢN TRỊ & QUẢN LÝ','Chiến lược, quản trị rủi ro, kiểm soát, khẩu vị rủi ro và chính sách tài chính.'),
        ('SO SÁNH NHÓM TƯƠNG ĐỒNG','Vị trí tương đối của doanh nghiệp so với trung bình/trung vị ngành.'),
        ('ANCHOR','Điểm khởi đầu theo đúng methodology áp dụng.'),
        ('ĐIỀU CHỈNH NỘI SINH / MODIFIERS','Từng notch/modifier, luận cứ và tác động.'),
        ('SACP / SCA','Năng lực tín nhiệm độc lập sau các điều chỉnh nội sinh.'),
        ('HỖ TRỢ BÊN NGOÀI','Khả năng và động cơ hỗ trợ từ tập đoàn/chính phủ khi phù hợp.'),
        ('ICR','Kết quả xếp hạng tín nhiệm tổ chức phát hành.'),
        ('ĐỘ NHẠY XẾP HẠNG','Điều kiện định lượng/định tính có thể dẫn tới nâng hoặc hạ hạng.'),
        ('STRESS TEST','Kịch bản bất lợi và tác động tới vốn, thanh khoản, lợi nhuận và rating.'),
        ('EARLY WARNING','Chỉ báo giám sát sau xếp hạng và ngưỡng cần đánh giá lại.'),
        ('KẾT LUẬN TRÌNH HỘI ĐỒNG XHTN','Tóm tắt waterfall, điểm tranh luận và vấn đề cần Hội đồng quyết định.'),
        ('PHỤ LỤC KPI','Bảng KPI chính và benchmark ngành.'),
        ('PHỤ LỤC WATERFALL','Audit trail từ Anchor tới ICR.'),
        ('PHỤ LỤC PEER & COVERAGE','Danh sách peer, trung bình/trung vị và số DN có dữ liệu.'),
        ('GIẢ ĐỊNH, HẠN CHẾ & TUYÊN BỐ SỬ DỤNG','Các giả định, override, giới hạn dữ liệu và điều kiện sử dụng.')
        ]
    return [
        ('TÓM TẮT ĐIỀU HÀNH','Luận điểm đầu tư, vùng giá và rủi ro trọng yếu.'),
        ('HỒ SƠ DOANH NGHIỆP','Loại hình, ngành, sàn, quy mô và vị trí trong peer group.'),
        ('BỐI CẢNH VĨ MÔ','Các biến vĩ mô có ảnh hưởng tới hoạt động và định giá.'),
        ('TRIỂN VỌNG NGÀNH','Chu kỳ, cạnh tranh, quy mô thị trường và động lực tăng trưởng.'),
        ('VỊ THẾ CẠNH TRANH','Franchise, thị phần, lợi thế và hạn chế.'),
        ('MÔ HÌNH KINH DOANH','Cơ cấu nguồn thu/tài sản và mức độ đa dạng.'),
        ('QUY MÔ & TĂNG TRƯỞNG','Tăng trưởng doanh thu/tài sản và chất lượng tăng trưởng.'),
        ('ROE & HIỆU QUẢ VỐN','ROE lịch sử, peer benchmark và nguồn tạo ROE.'),
        ('ROA & HIỆU QUẢ TÀI SẢN','ROA lịch sử và hiệu quả sử dụng tài sản.'),
        ('CHẤT LƯỢNG LỢI NHUẬN','Tính lặp lại, biên lợi nhuận và các khoản bất thường.'),
        ('CHẤT LƯỢNG TÀI SẢN','NPL hoặc rủi ro tài sản đặc thù.'),
        ('VỐN & ĐÒN BẨY','CAR/đòn bẩy và khả năng hấp thụ tổn thất.'),
        ('NGUỒN VỐN & THANH KHOẢN','Cơ cấu nguồn vốn, chi phí vốn và thanh khoản.'),
        ('DÒNG TIỀN','Dòng tiền hoạt động, FCF và khả năng tài trợ tăng trưởng.'),
        ('SO SÁNH PEER','Trung bình, trung vị ngành và vị trí tương đối.'),
        ('ĐỊNH GIÁ TƯƠNG ĐỐI','P/E, P/B, EV/EBITDA và benchmark phù hợp.'),
        ('ĐỊNH GIÁ NỘI TẠI / CƠ SỞ','Giá trị cơ sở và các giả định trọng yếu.'),
        ('BEAR - BASE - BULL','Vùng giá theo các kịch bản thận trọng, cơ sở và tích cực.'),
        ('ĐỘ NHẠY ĐỊNH GIÁ','Độ nhạy với ROE/COE/tăng trưởng hoặc biến số chuyên ngành.'),
        ('M&A - GIÁ TRỊ ĐỘC LẬP','Giá trị trước quyền kiểm soát và cộng hưởng.'),
        ('M&A - QUYỀN KIỂM SOÁT & SCARCITY','Control premium, scarcity và quy mô lô khi có cơ sở.'),
        ('M&A - CỘNG HƯỞNG','Cộng hưởng doanh thu, chi phí, vốn và nguồn vốn.'),
        ('PPA, GOODWILL & TÁI CẤU TRÚC','PPA, goodwill, tăng vốn, giảm nợ và xử lý tài sản.'),
        ('STRESS TEST','Kịch bản bất lợi và tác động tới tài chính/định giá.'),
        ('CATALYST & RỦI RO ĐẦU TƯ','Sự kiện mở khóa giá trị và yếu tố làm suy yếu luận điểm.'),
        ('KẾT LUẬN PHÂN TÍCH','Tổng hợp fundamental, peer, valuation và M&A.'),
        ('PHỤ LỤC KPI, PEER & BIỂU ĐỒ','KPI, benchmark và chuỗi thời gian.'),
        ('GIẢ ĐỊNH, DỮ LIỆU & TUYÊN BỐ SỬ DỤNG','Lineage, coverage, giả định và giới hạn.')
        ]

def _page_narrative(ticker,head,desc,meta,s,val,rr,report_type):
    a=intelligent_analyze(ticker)
    base=[desc]
    if report_type=='rating':
        base.append(f"Phương pháp áp dụng: {rr.get('MethodologyName','N/A')}. Anchor {rr.get('Anchor','N/A')}; SACP/SCA {rr.get('SACP',rr.get('SCA','N/A'))}; ICR {rr.get('ICR','N/A')}.")
        if head in ('SO SÁNH NHÓM TƯƠNG ĐỒNG','PHỤ LỤC PEER & COVERAGE'):
            base.append(f"Benchmark: {industry_label(ticker)}. Kết luận phải đọc cùng số doanh nghiệp có dữ liệu ở từng KPI và từng kỳ.")
        if head in ('ĐIỀU CHỈNH NỘI SINH / MODIFIERS','KẾT LUẬN TRÌNH HỘI ĐỒNG XHTN'):
            base.append("Mọi notch/modifier cần có bằng chứng định lượng, nhận định định tính và audit trail. Analyst Override phải ghi rõ lý do và không được làm mất dấu kết quả máy tính ban đầu.")
        if head in ('ĐỘ NHẠY XẾP HẠNG','STRESS TEST','EARLY WARNING'):
            base.append("Độ nhạy được trình bày theo điều kiện có thể kiểm chứng, liên kết trực tiếp với yếu tố có khả năng làm thay đổi notch, modifier, cap/floor hoặc mức hỗ trợ.")
    else:
        fv=fair_value_range(ticker)
        base.append(f"Giá trị mô hình: Bear {price(fv.get('Bear'))}; Base {price(fv.get('Base'))}; Bull {price(fv.get('Bull'))}; Strategic/M&A {price(fv.get('StrategicMA'))}.")
        if head in ('SO SÁNH PEER','PHỤ LỤC KPI, PEER & BIỂU ĐỒ'):
            base.append(f"Benchmark: {industry_label(ticker)}; sử dụng cả trung bình và trung vị ngành để giảm ảnh hưởng của outlier.")
        if head in ('BEAR - BASE - BULL','ĐỘ NHẠY ĐỊNH GIÁ','STRESS TEST'):
            base.append("Các kịch bản là kiểm tra độ nhạy, không phải dự báo chắc chắn. Các giả định phải được hiệu chỉnh theo chu kỳ ngành, chất lượng doanh nghiệp và dữ liệu mới nhất.")
        if head.startswith('M&A') or head.startswith('PPA'):
            base.append("Giá trị M&A là lớp giá trị riêng của từng giao dịch. Platform không áp dụng vùng giá 80.000-100.000 đồng/cp hoặc control premium của STB cho doanh nghiệp khác.")
    if head in ('TÓM TẮT ĐIỀU HÀNH','KẾT LUẬN PHÂN TÍCH','CATALYST & RỦI RO ĐẦU TƯ','TÓM TẮT XẾP HẠNG'):
        base.extend([x for x in [a.get('Conclusion')] if x])
        base.extend(a.get('Interpretations',[])[:2])
    # Chỉ đưa boilerplate kiểm soát vào các phần thật sự cần thiết.
    # Tránh lặp cùng một đoạn trên hàng chục trang của báo cáo.
    governance_heads={'TÓM TẮT XẾP HẠNG','PHẠM VI XẾP HẠNG & DỮ LIỆU','PHƯƠNG PHÁP LUẬN ÁP DỤNG',
                      'KẾT LUẬN TRÌNH HỘI ĐỒNG XHTN','GIẢ ĐỊNH, HẠN CHẾ & TUYÊN BỐ SỬ DỤNG'}
    if report_type=='rating' and head in governance_heads:
        base.append("Đánh giá tương đối phải chỉ ra doanh nghiệp tốt hơn, tương đương hay yếu hơn peer ở chỉ tiêu trọng yếu nào; mọi notch/modifier cần có bằng chứng và audit trail.")
        base.append("Khi dữ liệu định lượng và nhận định định tính mâu thuẫn, chuyên viên phải giải thích nguyên nhân thay vì để engine tự động ưu tiên một phía.")
    elif report_type!='rating' and head in {'TÓM TẮT ĐIỀU HÀNH','KẾT LUẬN PHÂN TÍCH','GIẢ ĐỊNH, DỮ LIỆU & TUYÊN BỐ SỬ DỤNG'}:
        base.append("Giá trị hợp lý được đọc như một vùng; Bear-Base-Bull thể hiện bất định của giả định, còn Strategic/M&A chỉ được ghi nhận khi quyền kiểm soát, scarcity hoặc synergy có thể chứng minh.")

    if head in {'TÓM TẮT XẾP HẠNG','PHẠM VI XẾP HẠNG & DỮ LIỆU','KẾT LUẬN TRÌNH HỘI ĐỒNG XHTN'} and report_type=='rating':
        ev=rating_evidence(ticker)
        base.append(f"Độ tin cậy XHTN tự động: {ev.get('RatingConfidence')}; độ đầy đủ dữ liệu cốt lõi {ev.get('DataQuality',{}).get('Coverage',0)*100:.0f}%.")
    elif head in {'TÓM TẮT ĐIỀU HÀNH','KẾT LUẬN PHÂN TÍCH'} and report_type!='rating':
        vt=triangulate(ticker)
        base.append(f"Độ tin cậy phân tích: {vt.get('AnalyticalConfidence')}; độ đầy đủ dữ liệu cốt lõi {vt.get('DataQuality',{}).get('Coverage',0)*100:.0f}%.")

    if head in governance_heads or head in {'TÓM TẮT ĐIỀU HÀNH','KẾT LUẬN PHÂN TÍCH'}:
        base.append("Nhận định cuối cùng phải được đối chiếu với BCTC hợp nhất, thuyết minh, công bố thông tin và dữ liệu định tính trước khi phát hành chính thức.")
    return base

def _evidence_table(doc,ticker,head,report_type):
    if head in ('SO SÁNH PEER','SO SÁNH NHÓM TƯƠNG ĐỒNG','PHỤ LỤC PEER & COVERAGE'):
        return
    k,_,_=sector_kpi_table(ticker)
    if not len(k):return
    meta=get_company(ticker)
    wanted=_page_peer_metrics(head,meta.get('EntityType'))
    # Avoid repeating the same five rows on consecutive pages.
    if not wanted:return
    kk=k[k['Metric'].astype(str).isin(wanted)].drop_duplicates('Metric').head(6)
    if not len(kk):return
    t=doc.add_table(rows=1,cols=5);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=['Chỉ tiêu','Doanh nghiệp','TB ngành','Trung vị','Số DN']
    for j,x in enumerate(hdr):t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),'E8EEF7')
    for _,r in kk.iterrows():
        c=t.add_row().cells;m=r['Metric']
        vals=[r['Chỉ tiêu'],metric_fmt(m,r['Doanh nghiệp']),metric_fmt(m,r['Trung bình ngành']),
              metric_fmt(m,r['Trung vị ngành']),str(int(r['Số DN có dữ liệu']))]
        for j,x in enumerate(vals):c[j].text=str(x)
    for row in t.rows:
        for cell in row.cells:
            _set_cell_margins(cell,top=35,start=55,bottom=35,end=55)
            for pp in cell.paragraphs:
                pp.paragraph_format.space_after=Pt(0);pp.paragraph_format.line_spacing=1.0
                for run in pp.runs:run.font.name='Lato';run.font.size=Pt(8.5)

def _add_evidence_ledger(doc,ticker):
    ev=rating_evidence(ticker); rows=ev.get('EvidenceLedger',[])
    t=doc.add_table(rows=1,cols=5);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=['Cấu phần','Kết quả','Điều chỉnh','Bằng chứng','Kiểm chứng']
    for j,x in enumerate(hdr):t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),'E8EEF7')
    for r in rows:
        c=t.add_row().cells
        vals=[r.get('Cấu phần'),r.get('Kết quả'),r.get('Điều chỉnh'),r.get('Loại bằng chứng'),r.get('Trạng thái kiểm chứng')]
        for j,x in enumerate(vals):c[j].text=str(x)
    for row in t.rows:
        for cell in row.cells:
            _set_cell_margins(cell)
            for pp in cell.paragraphs:
                for run in pp.runs:run.font.name='Lato';run.font.size=Pt(9)

def _add_triangulation_table(doc,ticker):
    vt=triangulate(ticker);rows=vt.get('Lenses',[])
    t=doc.add_table(rows=1,cols=3);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=['Lăng kính','Giá trị','Trạng thái']
    for j,x in enumerate(hdr):t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),'E8EEF7')
    for r in rows:
        c=t.add_row().cells
        c[0].text=str(r.get('Lăng kính')); c[1].text=price(r.get('Giá trị')); c[2].text=str(r.get('Trạng thái'))
    for row in t.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                for run in pp.runs:run.font.name='Lato';run.font.size=Pt(10)


def peer_bar_chart(ticker,metric,title=None,top_n=12):
    from scripts.universal_data import industry_snapshot
    peers=industry_snapshot(ticker)
    if peers is None or not len(peers) or metric not in peers.columns:return None
    z=peers[['Ticker',metric]].copy()
    z[metric]=pd.to_numeric(z[metric],errors='coerce')
    z=z.dropna(subset=[metric]).drop_duplicates('Ticker')
    if not len(z):return None
    selected=str(ticker).upper()
    company=z[z.Ticker.astype(str).str.upper().eq(selected)]
    if len(company):
        cv=float(company.iloc[-1][metric]); z['dist']=(z[metric]-cv).abs()
        chosen=pd.concat([company,z[~z.Ticker.astype(str).str.upper().eq(selected)].sort_values('dist').head(max(0,top_n-1))])
    else:
        chosen=z.sort_values(metric).tail(top_n)
    chosen=chosen.drop_duplicates('Ticker').sort_values(metric)
    mean=float(z[metric].mean())
    fig,ax=plt.subplots(figsize=(8.4,4.2));plt.rcParams.update({'font.family':'Lato','font.size':10})
    ax.barh(chosen.Ticker.astype(str),chosen[metric]);ax.axvline(mean,linestyle='--',linewidth=1.5,label='Trung bình ngành')
    ax.set_title(title or f"{VI_METRIC.get(metric,metric)} - so sánh peer",fontsize=11);ax.grid(axis='x',alpha=.2);ax.legend(fontsize=8)
    if metric in PCT:ax.xaxis.set_major_formatter(lambda v,pos:(f'{v*100:.1f}%').replace('.',','))
    fig.tight_layout();bio=BytesIO();fig.savefig(bio,dpi=180,bbox_inches='tight');plt.close(fig);bio.seek(0);return bio

def peer_scatter_chart(ticker,xmetric,ymetric,title=None):
    from scripts.universal_data import industry_snapshot
    peers=industry_snapshot(ticker)
    if peers is None or not len(peers) or xmetric not in peers.columns or ymetric not in peers.columns:return None
    z=peers[['Ticker',xmetric,ymetric]].copy()
    z[xmetric]=pd.to_numeric(z[xmetric],errors='coerce');z[ymetric]=pd.to_numeric(z[ymetric],errors='coerce')
    z=z.dropna(subset=[xmetric,ymetric]).drop_duplicates('Ticker')
    if len(z)<2:return None
    fig,ax=plt.subplots(figsize=(8.4,4.4));plt.rcParams.update({'font.family':'Lato','font.size':10})
    ax.scatter(z[xmetric],z[ymetric],alpha=.7)
    for _,r in z.iterrows():
        if str(r.Ticker).upper()==str(ticker).upper():
            ax.annotate(str(r.Ticker),(r[xmetric],r[ymetric]),xytext=(5,5),textcoords='offset points',fontweight='bold')
    ax.axvline(z[xmetric].mean(),linestyle='--',linewidth=1);ax.axhline(z[ymetric].mean(),linestyle='--',linewidth=1)
    ax.set_xlabel(VI_METRIC.get(xmetric,xmetric));ax.set_ylabel(VI_METRIC.get(ymetric,ymetric))
    ax.set_title(title or f"{VI_METRIC.get(ymetric,ymetric)} và {VI_METRIC.get(xmetric,xmetric)} - vị trí tương đối",fontsize=11)
    if xmetric in PCT:ax.xaxis.set_major_formatter(lambda v,pos:(f'{v*100:.1f}%').replace('.',','))
    if ymetric in PCT:ax.yaxis.set_major_formatter(lambda v,pos:(f'{v*100:.1f}%').replace('.',','))
    ax.grid(alpha=.2);fig.tight_layout();bio=BytesIO();fig.savefig(bio,dpi=180,bbox_inches='tight');plt.close(fig);bio.seek(0);return bio

def _relative_analysis(ticker,metric):
    k,_,_=sector_kpi_table(ticker)
    if not len(k):return None
    r=k[k.Metric.astype(str).eq(str(metric))]
    if not len(r):return None
    r=r.iloc[-1]
    c=num(r['Doanh nghiệp']);m=num(r['Trung bình ngành']);med=num(r['Trung vị ngành']);n=int(r['Số DN có dữ liệu'])
    if c is None:return None
    label=r['Chỉ tiêu']
    if m in (None,0):return f"{label} của doanh nghiệp là {metric_fmt(metric,c)}; chưa đủ dữ liệu ngành để kết luận tương đối."
    gap=c/m-1;direction='cao hơn' if gap>0 else 'thấp hơn';mag=abs(gap);g=(f"{mag*100:.1f}").replace('.',',')
    txt=f"{label} đạt {metric_fmt(metric,c)}, {direction} trung bình ngành {metric_fmt(metric,m)} khoảng {g}% (mẫu {n} doanh nghiệp có dữ liệu)."
    if med is not None:txt+=f" Trung vị ngành là {metric_fmt(metric,med)}, dùng để kiểm tra ảnh hưởng của outlier."
    if metric=='ROE':txt+=" ROE cần đọc cùng đòn bẩy, chất lượng lợi nhuận và mức định giá P/B/P/E."
    elif metric=='NPL':txt+=" Chênh lệch NPL cần đọc cùng mức bao phủ dự phòng, tài sản bảo đảm và bộ đệm vốn."
    elif metric=='CAR':txt+=" CAR cần đọc cùng tăng trưởng tài sản có rủi ro, kế hoạch vốn và chất lượng tài sản."
    elif metric=='CASA':txt+=" CASA cao hơn peer có thể hỗ trợ chi phí vốn/NIM nhưng cần kiểm tra độ ổn định tiền gửi."
    elif metric=='DebtEquity':txt+=" Đòn bẩy cao hơn peer làm tăng độ nhạy lợi nhuận và khả năng trả nợ."
    return txt

def _page_peer_metrics(head,entity_type):
    if entity_type=='BANK':
        mp={
        'QUY MÔ & TĂNG TRƯỞNG':['TotalAssets','GrossLoans','CustomerDeposits'],
        'ROE & HIỆU QUẢ VỐN':['ROE','NIM','CIR'],
        'ROA & HIỆU QUẢ TÀI SẢN':['ROA','LoanAssets','EquityAssets'],
        'KHẢ NĂNG SINH LỢI':['ROE','ROA','NIM','CIR'],
        'CHẤT LƯỢNG TÀI SẢN':['NPL','ProvisionOperatingIncome','LoanAssets'],
        'CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN':['NPL','ProvisionOperatingIncome','LoanAssets'],
        'VỐN & ĐÒN BẨY':['CAR','EquityAssets','ROE'],
        'NGUỒN VỐN & THANH KHOẢN':['CASA','LDR','DepositAssets'],
        'THANH KHOẢN':['LDR','CASA','DepositAssets'],
        'SO SÁNH PEER':['ROE','ROA','NIM','CIR','NPL','CAR','CASA','LDR','PB'],
        'SO SÁNH NHÓM TƯƠNG ĐỒNG':['ROE','ROA','NIM','CIR','NPL','CAR','CASA','LDR','PB']}
    elif entity_type=='SECURITIES':
        mp={
        'QUY MÔ & TĂNG TRƯỞNG':['Revenue','TotalAssets','Equity'],
        'ROE & HIỆU QUẢ VỐN':['ROE','AvailableCapitalRatio','DebtEquity'],
        'ROA & HIỆU QUẢ TÀI SẢN':['ROA','NetMargin'],
        'KHẢ NĂNG SINH LỢI':['ROE','ROA','NetMargin','ICGR'],
        'CHẤT LƯỢNG TÀI SẢN':['MarginLoansEquity','DebtEquity'],
        'CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN':['MarginLoansEquity','DebtEquity'],
        'VỐN & ĐÒN BẨY':['AvailableCapitalRatio','DebtEquity','DebtEBITDA'],
        'NGUỒN VỐN & THANH KHOẢN':['CurrentRatio','DebtEquity','CFO_Debt'],
        'THANH KHOẢN':['CurrentRatio','CFO_Debt'],
        'SO SÁNH PEER':['ROE','ROA','AvailableCapitalRatio','DebtEquity','CurrentRatio','PB','PE'],
        'SO SÁNH NHÓM TƯƠNG ĐỒNG':['ROE','ROA','AvailableCapitalRatio','DebtEquity','CurrentRatio','PB','PE']}
    else:
        mp={
        'QUY MÔ & TĂNG TRƯỞNG':['Revenue','TotalAssets','GrossMargin'],
        'ROE & HIỆU QUẢ VỐN':['ROE','DebtEquity','NetMargin'],
        'ROA & HIỆU QUẢ TÀI SẢN':['ROA','EBITDAMargin'],
        'KHẢ NĂNG SINH LỢI':['ROE','ROA','GrossMargin','NetMargin','EBITDAMargin'],
        'CHẤT LƯỢNG TÀI SẢN':['CFO_Debt','FOCF_Debt','DebtEBITDA'],
        'CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN':['CFO_Debt','FOCF_Debt','DebtEBITDA'],
        'VỐN & ĐÒN BẨY':['DebtEquity','DebtEBITDA','CFO_Debt','FOCF_Debt'],
        'NGUỒN VỐN & THANH KHOẢN':['CurrentRatio','CFO_Debt','FOCF_Debt'],
        'THANH KHOẢN':['CurrentRatio','CFO_Debt','FOCF_Debt'],
        'SO SÁNH PEER':['ROE','ROA','GrossMargin','NetMargin','DebtEquity','DebtEBITDA','CurrentRatio','PB','PE'],
        'SO SÁNH NHÓM TƯƠNG ĐỒNG':['ROE','ROA','GrossMargin','NetMargin','DebtEquity','DebtEBITDA','CurrentRatio','PB','PE']}
    if head in ('ĐỊNH GIÁ TƯƠNG ĐỐI','ĐỊNH GIÁ CƠ SỞ & CHẾ ĐỘ ĐỊNH GIÁ'):return ['PB','PE','EV_EBITDA']
    return mp.get(head,[])

def _add_peer_section(doc,ticker,head,meta):
    """Body report: chỉ giữ tối đa 1 đồ thị peer đại diện cho mỗi phần.
    Phần phụ lục vẫn chứa bộ đồ thị peer đầy đủ. Cách này tránh lặp narrative/biểu đồ
    và giúp Word dồn nội dung liên tục thay vì để khoảng trắng lớn quanh các ảnh cao.
    """
    metrics=_page_peer_metrics(head,meta.get('EntityType'))
    if not metrics:return
    mm=metrics[0]
    try:
        bio=peer_bar_chart(ticker,mm)
        if bio:doc.add_picture(bio,width=Mm(160))
    except Exception:pass
    if head in ('SO SÁNH PEER','SO SÁNH NHÓM TƯƠNG ĐỒNG'):
        try:
            bio=peer_scatter_chart(ticker,'ROE','PB')
            if bio:doc.add_picture(bio,width=Mm(160))
        except Exception:pass


def _decision_narrative(ticker,head,meta,snapshot,val,report_type):
    """Four-question analyst narrative: What? Why? Peer? So what?
    Uses only model/public-data fields already available; avoids inventing causes.
    """
    et=meta.get('EntityType')
    lines=[]
    metrics=_page_peer_metrics(head,et)
    # 1) What do the numbers say + peer comparison
    if metrics:
        facts=[]
        for m in metrics[:3]:
            try:
                x=_relative_analysis(ticker,m)
                if x:facts.append(x)
            except Exception:pass
        if facts:
            lines.append("Số liệu nói gì? "+" ".join(facts))
    # 2) Why: disciplined driver framing, not unsupported causal claims
    driver_map={
      'ROE & HIỆU QUẢ VỐN':"Động lực cần kiểm chứng gồm biên lợi nhuận, hiệu suất sử dụng tài sản và đòn bẩy. Báo cáo không quy kết nguyên nhân khi dữ liệu nguồn chưa đủ bằng chứng.",
      'ROA & HIỆU QUẢ TÀI SẢN':"ROA phản ánh khả năng chuyển quy mô tài sản thành lợi nhuận; cần đối chiếu tăng trưởng tài sản, biên lợi nhuận và các khoản thu nhập bất thường.",
      'KHẢ NĂNG SINH LỢI':"Khả năng sinh lợi bền vững cần được kiểm tra qua biên lợi nhuận, cơ cấu thu nhập, chi phí hoạt động và mức sử dụng đòn bẩy.",
      'CHẤT LƯỢNG TÀI SẢN':"Chất lượng tài sản cần được kiểm tra đồng thời qua nợ xấu, dự phòng, tăng trưởng tín dụng/tài sản và mức tập trung rủi ro.",
      'CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN':"Chất lượng tài sản cần được kiểm tra đồng thời qua nợ xấu, dự phòng, tăng trưởng tín dụng/tài sản và mức tập trung rủi ro.",
      'VỐN & ĐÒN BẨY':"Bộ đệm vốn phải được đặt cạnh tốc độ tăng trưởng, chất lượng tài sản và khả năng tạo vốn nội bộ; một tỷ lệ vốn cao không tự động đồng nghĩa rủi ro thấp.",
      'NGUỒN VỐN & THANH KHOẢN':"Chất lượng nguồn vốn phụ thuộc chi phí, độ ổn định, mức tập trung và khả năng chuyển đổi tài sản thành thanh khoản trong điều kiện căng thẳng.",
      'THANH KHOẢN':"Thanh khoản cần được đọc theo cả trạng thái hiện tại và sức chịu đựng khi dòng tiền bất lợi; tỷ lệ kế toán đơn lẻ không đủ để kết luận.",
      'ĐỊNH GIÁ TƯƠNG ĐỐI':"Premium/discount định giá chỉ hợp lý khi tương xứng với ROE, tăng trưởng, rủi ro và chất lượng lợi nhuận so với peer.",
      'ĐỊNH GIÁ CƠ SỞ & CHẾ ĐỘ ĐỊNH GIÁ':"Giá trị hợp lý được xem như một vùng thay vì một điểm duy nhất; độ nhạy với ROE, COE, tăng trưởng và multiple giúp lượng hóa rủi ro sai số.",
      'SO SÁNH PEER':"So sánh peer ưu tiên cả trung bình và trung vị để giảm ảnh hưởng của outlier; vị trí tương đối quan trọng hơn một ngưỡng tuyệt đối.",
      'SO SÁNH NHÓM TƯƠNG ĐỒNG':"So sánh peer ưu tiên cả trung bình và trung vị để giảm ảnh hưởng của outlier; vị trí tương đối quan trọng hơn một ngưỡng tuyệt đối."
    }
    if head in driver_map:lines.append("Tại sao cần chú ý? "+driver_map[head])

    # 3/4) Explicit implication for valuation vs credit rating
    if report_type=='analysis':
        implication={
          'ROE & HIỆU QUẢ VỐN':"Tác động đến giá cổ phiếu: ROE cao và bền vững hơn peer có thể biện minh cho P/B premium; ngược lại, ROE thấp nhưng P/B cao làm tăng rủi ro định giá.",
          'ROA & HIỆU QUẢ TÀI SẢN':"Tác động đến giá cổ phiếu: ROA tốt hơn peer hỗ trợ chất lượng lợi nhuận và khả năng duy trì ROE mà không cần tăng mạnh đòn bẩy.",
          'KHẢ NĂNG SINH LỢI':"Tác động đến giá cổ phiếu: chất lượng và độ bền của lợi nhuận quyết định mức multiple có thể duy trì, không chỉ tốc độ tăng EPS ngắn hạn.",
          'CHẤT LƯỢNG TÀI SẢN':"Tác động đến giá cổ phiếu: suy giảm chất lượng tài sản có thể làm tăng chi phí tín dụng, giảm ROE kỳ vọng và kéo giảm multiple hợp lý.",
          'VỐN & ĐÒN BẨY':"Tác động đến giá cổ phiếu: bộ đệm vốn tốt tạo dư địa tăng trưởng/phân phối vốn; thiếu vốn có thể dẫn đến pha loãng hoặc hạn chế tăng trưởng.",
          'NGUỒN VỐN & THANH KHOẢN':"Tác động đến giá cổ phiếu: nguồn vốn ổn định và chi phí thấp hỗ trợ biên lợi nhuận, đồng thời giảm tail risk thanh khoản.",
          'ĐỊNH GIÁ TƯƠNG ĐỐI':"Kết luận định giá phải trả lời premium/discount so với peer có được giải thích bởi ROE, tăng trưởng và rủi ro hay không.",
          'SO SÁNH PEER':"Kết luận đầu tư: ưu tiên doanh nghiệp có tổ hợp ROE/tăng trưởng/chất lượng tài sản tốt hơn peer nhưng valuation chưa phản ánh đầy đủ lợi thế đó.",
          'SO SÁNH NHÓM TƯƠNG ĐỒNG':"Kết luận đầu tư: ưu tiên doanh nghiệp có tổ hợp ROE/tăng trưởng/chất lượng tài sản tốt hơn peer nhưng valuation chưa phản ánh đầy đủ lợi thế đó."
        }
    else:
        implication={
          'ROE & HIỆU QUẢ VỐN':"Tác động đến XHTN: khả năng sinh lợi tốt tạo vốn nội bộ và hấp thụ tổn thất; nhưng lợi nhuận dựa nhiều vào đòn bẩy hoặc thu nhập bất thường có chất lượng thấp hơn.",
          'ROA & HIỆU QUẢ TÀI SẢN':"Tác động đến XHTN: hiệu quả tài sản tốt hỗ trợ khả năng tạo bộ đệm vốn, song phải được kiểm tra tính bền vững.",
          'KHẢ NĂNG SINH LỢI':"Tác động đến XHTN: lợi nhuận bền vững củng cố năng lực hấp thụ tổn thất và trả nợ; biến động lợi nhuận làm giảm độ chắc chắn của hồ sơ tài chính.",
          'CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN':"Tác động đến XHTN: chất lượng tài sản suy yếu có thể truyền dẫn sang dự phòng, lợi nhuận, vốn và thanh khoản, do đó là biến số trọng yếu của notch.",
          'VỐN & ĐÒN BẨY':"Tác động đến XHTN: bộ đệm vốn/đòn bẩy quyết định khả năng hấp thụ tổn thất ngoài dự kiến và là đầu vào quan trọng cho đánh giá hồ sơ tài chính.",
          'THANH KHOẢN':"Tác động đến XHTN: thanh khoản yếu có thể tạo áp lực trả nợ ngay cả khi doanh nghiệp còn khả năng sinh lợi; cần xem xét cùng cấu trúc đáo hạn và nguồn dự phòng.",
          'SO SÁNH NHÓM TƯƠNG ĐỒNG':"Tác động đến XHTN: peer comparison là phép kiểm tra tính hợp lý của đánh giá định tính và notch, không thay thế methodology."
        }
    if head in implication:lines.append(implication[head])
    return lines


def _fmt_method_value(metric,v):
    try:v=float(v)
    except:return 'N/A'
    if not np.isfinite(v):return 'N/A'
    if metric in METH_PCT:return f"{v*100:.1f}%".replace('.',',')
    if metric in METH_MULT:return f"{v:.2f}x".replace('.',',')
    if abs(v)>=1e12:return (f"{v/1e12:,.1f} nghìn tỷ").replace(',','X').replace('.',',').replace('X','.')
    if abs(v)>=1e9:return (f"{v/1e9:,.1f} tỷ").replace(',','X').replace('.',',').replace('X','.')
    return (f"{v:,.2f}").replace(',','X').replace('.',',').replace('X','.')

def _add_methodology_kpi_matrix(doc,ticker):
    z=methodology_kpi_table(ticker,include_missing=True)
    if z.empty:return
    doc.add_heading('MA TRẬN CHỈ TIÊU THEO PHƯƠNG PHÁP XHTN',1)
    _add_pars(doc,["Ma trận dưới đây mở rộng phạm vi phân tích theo đúng loại hình doanh nghiệp. Chỉ tiêu chưa có dữ liệu được giữ N/A để chỉ rõ data gap, không tự giả định số liệu."])
    for grp,g in z.groupby('Nhóm phân tích',sort=False):
        doc.add_heading(str(grp),2)
        tb=doc.add_table(rows=1,cols=6);tb.style='Table Grid'
        hdr=['Chỉ tiêu','Doanh nghiệp','TB ngành','Trung vị','Số DN','Dữ liệu']
        for j,x in enumerate(hdr):tb.rows[0].cells[j].text=x
        for _,r in g.iterrows():
            c=tb.add_row().cells;m=r['Metric']
            vals=[r['Chỉ tiêu'],_fmt_method_value(m,r['Doanh nghiệp']),_fmt_method_value(m,r['TB ngành']),
                  _fmt_method_value(m,r['Trung vị']),str(int(r['Số DN'])),r['Trạng thái dữ liệu']]
            for j,x in enumerate(vals):c[j].text=str(x)


GREEN='4F9E1E'
LIGHT_GREEN='E6F3D8'
DARK_GREEN='2F6F12'
GREY='F2F2F2'


def _set_cell_border(cell, **edges):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge not in edges: continue
        tag='w:'+edge; el=borders.find(qn(tag))
        if el is None: el=OxmlElement(tag); borders.append(el)
        for k,v in edges[edge].items(): el.set(qn('w:'+k),str(v))


def _section_band(doc,text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); _set_cell_shading(c,GREEN); _set_cell_margins(c,top=60,start=90,bottom=60,end=90)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.keep_with_next=True
    r=p.add_run(str(text).upper()); r.bold=True; r.font.name='Lato'; r.font.size=Pt(12); r.font.color.rgb=None
    # white text through OOXML for robust LO rendering
    color=OxmlElement('w:color'); color.set(qn('w:val'),'FFFFFF'); r._r.get_or_add_rPr().append(color)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def _subhead(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2); p.paragraph_format.keep_with_next=True
    r=p.add_run(text); r.bold=True; r.font.name='Lato'; r.font.size=Pt(11); color=OxmlElement('w:color'); color.set(qn('w:val'),DARK_GREEN); r._r.get_or_add_rPr().append(color)
    return p



def _intel_body(doc, text):
    """Consistent public-intelligence body style: same Lato face/size/spacing as analytical narrative."""
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing=1.05
    p.paragraph_format.space_after=Pt(3)
    r=p.add_run(str(text))
    r.font.name='Lato'; r.font.size=Pt(9.2)
    return p

def _intel_source(doc, source, asof, url=None):
    """Compact source note; keep URL in data/audit trail instead of printing a long raw URL in the report."""
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0)
    p.paragraph_format.space_after=Pt(4)
    r=p.add_run(f"Nguồn: {source}; cập nhật {asof}.")
    r.font.name='Lato'; r.font.size=Pt(7.4); r.font.italic=True
    return p

def _compact_pars(doc,pars,max_pars=4):
    seen=set(); n=0
    for txt in pars:
        key=' '.join(str(txt).split()).strip()
        if not key or key.lower() in seen: continue
        seen.add(key.lower()); n+=1
        if n>max_pars: break
        p=doc.add_paragraph(key); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing=1.08; p.paragraph_format.space_after=Pt(3); p.paragraph_format.widow_control=True


def _add_two_charts(doc,items,width_mm=76):
    valid=[]
    for label,bio in items:
        if bio: valid.append((label,bio))
    if not valid:return
    tbl=doc.add_table(rows=1,cols=min(2,len(valid)));tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,(label,bio) in enumerate(valid[:2]):
        c=tbl.cell(0,j); _set_cell_margins(c,top=25,start=25,bottom=25,end=25)
        p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(1)
        r=p.add_run(label);r.bold=True;r.font.name='Lato';r.font.size=Pt(8.5)
        q=c.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_after=Pt(0)
        q.add_run().add_picture(bio,width=Mm(width_mm))
        _set_cell_border(c,bottom={'val':'single','sz':'4','color':'D9D9D9'})


def _cover_sample(doc,ticker,meta,report_type):
    # Dense, simple cover inspired by supplied reports: no metadata table that consumes space.
    for _ in range(3): doc.add_paragraph()
    t=doc.add_table(rows=1,cols=1);c=t.cell(0,0);_set_cell_shading(c,GREEN);_set_cell_margins(c,top=850,start=320,bottom=850,end=320)
    p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    title='BÁO CÁO XẾP HẠNG TÍN NHIỆM' if report_type=='rating' else 'BÁO CÁO PHÂN TÍCH CỔ PHIẾU'
    r=p.add_run(title+'\n');r.bold=True;r.font.name='Lato';r.font.size=Pt(30)
    col=OxmlElement('w:color');col.set(qn('w:val'),'FFFFFF');r._r.get_or_add_rPr().append(col)
    r=p.add_run(company_display_name(meta,ticker));r.bold=True;r.font.name='Lato';r.font.size=Pt(22)
    col=OxmlElement('w:color');col.set(qn('w:val'),'FFFFFF');r._r.get_or_add_rPr().append(col)
    p=doc.add_paragraph();p.paragraph_format.space_before=Pt(16);p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent=Mm(0);p.paragraph_format.right_indent=Mm(0)
    for i,(label,value) in enumerate([('NGÀNH',str(meta.get('Sector','')).upper()),('QUỐC GIA','VIỆT NAM')]):
        if i: p.add_run('\n')
        r=p.add_run(f"{label}: ");r.font.name='Lato';r.font.size=Pt(11);r.bold=True
        r=p.add_run(str(value));r.font.name='Lato';r.font.size=Pt(11);r.bold=True
    doc.add_page_break()


def _rating_summary_page(doc,ticker,meta,rr):
    # Header band + two-column rating card / thesis, matching structure of KLB/VDS samples.
    t=doc.add_table(rows=1,cols=1);c=t.cell(0,0);_set_cell_shading(c,GREEN);_set_cell_margins(c,top=100,start=110,bottom=100,end=110)
    p=c.paragraphs[0];r=p.add_run(company_display_name(meta,ticker).upper());r.bold=True;r.font.name='Lato';r.font.size=Pt(16)
    col=OxmlElement('w:color');col.set(qn('w:val'),'FFFFFF');r._r.get_or_add_rPr().append(col)
    p=c.add_paragraph(f"NGÀNH: {meta.get('Sector')}    |    PHƯƠNG PHÁP: {rr.get('MethodologyName',meta.get('Methodology','N/A'))}");p.paragraph_format.space_after=Pt(0)
    for run in p.runs:
        run.font.name='Lato';run.font.size=Pt(9);col=OxmlElement('w:color');col.set(qn('w:val'),'FFFFFF');run._r.get_or_add_rPr().append(col)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    tbl=doc.add_table(rows=1,cols=2);tbl.autofit=False;tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    left,right=tbl.cell(0,0),tbl.cell(0,1);left.width=Mm(57);right.width=Mm(112)
    _set_cell_shading(left,LIGHT_GREEN);_set_cell_margins(left,top=80,start=80,bottom=80,end=80);_set_cell_margins(right,top=30,start=100,bottom=30,end=70)
    p=left.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('KẾT QUẢ XẾP HẠNG');r.bold=True;r.font.size=Pt(10);r.font.name='Lato'
    rows=[('Bậc xếp hạng',rr.get('ICR','N/A')),('Triển vọng',rr.get('Outlook','Ổn định')),('Anchor',rr.get('Anchor','N/A')),('SACP / SCA',rr.get('SACP',rr.get('SCA','N/A'))),('Hỗ trợ bên ngoài',rr.get('ExternalSupport','Trung lập')),('Loại hình',entity_type_vi(meta.get('EntityType'))),('Benchmark',industry_label(ticker))]
    for a,b in rows:
        p=left.add_paragraph();p.paragraph_format.space_after=Pt(1);p.paragraph_format.line_spacing=1.0
        r=p.add_run(a+': ');r.bold=True;r.font.name='Lato';r.font.size=Pt(8.5)
        r=p.add_run(str(b));r.font.name='Lato';r.font.size=Pt(8.5)
    p=right.paragraphs[0];r=p.add_run('LUẬN ĐIỂM XẾP HẠNG');r.bold=True;r.font.name='Lato';r.font.size=Pt(11)
    a=intelligent_analyze(ticker)
    thesis=[a.get('Conclusion','')]
    if a.get('Strengths'): thesis.append('Điểm mạnh: '+' '.join(a['Strengths'][:2]))
    if a.get('Risks'): thesis.append('Điểm cần theo dõi: '+' '.join(a['Risks'][:2]))
    thesis.append(f"Kết quả mô hình hiện tại là {rr.get('ICR','N/A')}; đánh giá cuối cùng cần đối chiếu dữ liệu nguồn, peer và các yếu tố định tính trọng yếu.")
    for txt in thesis:
        if txt:
            p=right.add_paragraph(str(txt));p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.line_spacing=1.05;p.paragraph_format.space_after=Pt(4)
            for run in p.runs:run.font.name='Lato';run.font.size=Pt(9.5)
    _set_cell_border(left,right={'val':'single','sz':'6','color':'FFFFFF'})


def _analysis_summary_page(doc,ticker,meta,s,val):
    _section_band(doc,'TÓM TẮT')
    tbl=doc.add_table(rows=1,cols=2);tbl.autofit=False;tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    left,right=tbl.cell(0,0),tbl.cell(0,1);left.width=Mm(58);right.width=Mm(112)
    _set_cell_shading(left,GREY);_set_cell_margins(left,top=60,start=70,bottom=60,end=70);_set_cell_margins(right,top=20,start=100,bottom=20,end=50)
    basics=[('Mã',ticker),('P/B',mult(s.get('PB'))),('P/E',mult(s.get('PE'))),('ROE',pct(s.get('ROE'))),('ROA',pct(s.get('ROA'))),('Giá cơ sở',price(val.get('FairValue',val.get('BaseValue'))))]
    p=left.paragraphs[0];r=p.add_run('THÔNG TIN CƠ BẢN');r.bold=True;r.font.name='Lato';r.font.size=Pt(10)
    for a,b in basics:
        p=left.add_paragraph();p.paragraph_format.space_after=Pt(1);r=p.add_run(a+': ');r.bold=True;r.font.size=Pt(8.5);r.font.name='Lato';r=p.add_run(str(b));r.font.size=Pt(8.5);r.font.name='Lato'
    p=right.paragraphs[0];r=p.add_run('LUẬN ĐIỂM ĐẦU TƯ');r.bold=True;r.font.name='Lato';r.font.size=Pt(11)
    _compact_pars_cell=right
    a=intelligent_analyze(ticker)
    texts=[a.get('Conclusion','')]
    if a.get('Strengths'):texts.append('Luận điểm tích cực: '+' '.join(a['Strengths'][:3]))
    if a.get('Risks'):texts.append('Rủi ro chính: '+' '.join(a['Risks'][:3]))
    fv=fair_value_range(ticker);texts.append(f"Vùng giá: Bear {price(fv.get('Bear'))}; Base {price(fv.get('Base'))}; Bull {price(fv.get('Bull'))}.")
    for txt in texts:
        if txt:
            p=right.add_paragraph(str(txt));p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.line_spacing=1.05;p.paragraph_format.space_after=Pt(4)
            for run in p.runs:run.font.name='Lato';run.font.size=Pt(9.5)


def _rating_groups(entity_type):
    return [
      ('NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG',['HỒ SƠ KINH DOANH','KHẢ NĂNG SINH LỢI','CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN','VỐN & ĐÒN BẨY','THANH KHOẢN']),
      ('THÔNG TIN TỔNG QUAN TỔ CHỨC PHÁT HÀNH',['PHẠM VI XẾP HẠNG & DỮ LIỆU','QUY MÔ & TĂNG TRƯỞNG']),
      ('RỦI RO VĨ MÔ VÀ NGÀNH',['RỦI RO VĨ MÔ','RỦI RO NGÀNH']),
      ('HỒ SƠ KINH DOANH',['HỒ SƠ KINH DOANH','QUY MÔ & TĂNG TRƯỞNG','SO SÁNH NHÓM TƯƠNG ĐỒNG']),
      ('HỒ SƠ TÀI CHÍNH',['KHẢ NĂNG SINH LỢI','CHẤT LƯỢNG TÀI SẢN / RỦI RO TÀI SẢN','VỐN & ĐÒN BẨY','NGUỒN VỐN','THANH KHOẢN','DÒNG TIỀN & KHẢ NĂNG TRẢ NỢ']),
      ('QUẢN TRỊ VÀ QUẢN LÝ',['QUẢN TRỊ & QUẢN LÝ']),
      ('KHUNG XẾP HẠNG VÀ KẾT QUẢ',['PHƯƠNG PHÁP LUẬN ÁP DỤNG','ANCHOR','ĐIỀU CHỈNH NỘI SINH / MODIFIERS','SACP / SCA','HỖ TRỢ BÊN NGOÀI','ICR','ĐỘ NHẠY XẾP HẠNG','STRESS TEST','EARLY WARNING','KẾT LUẬN TRÌNH HỘI ĐỒNG XHTN']),
    ]


def _analysis_groups(entity_type):
    return [
      ('I. TỔNG QUAN',['HỒ SƠ DOANH NGHIỆP','VỊ THẾ CẠNH TRANH','MÔ HÌNH KINH DOANH']),
      ('II. HOẠT ĐỘNG KINH DOANH',['QUY MÔ & TĂNG TRƯỞNG','ROE & HIỆU QUẢ VỐN','ROA & HIỆU QUẢ TÀI SẢN','CHẤT LƯỢNG LỢI NHUẬN']),
      ('III. TÌNH HÌNH TÀI CHÍNH',['CHẤT LƯỢNG TÀI SẢN','VỐN & ĐÒN BẨY','NGUỒN VỐN & THANH KHOẢN','DÒNG TIỀN','SO SÁNH PEER']),
      ('IV. LUẬN ĐIỂM ĐẦU TƯ',['CATALYST & RỦI RO ĐẦU TƯ','M&A - GIÁ TRỊ ĐỘC LẬP','M&A - QUYỀN KIỂM SOÁT & SCARCITY','M&A - CỘNG HƯỞNG']),
      ('V. RỦI RO',['STRESS TEST','ĐỘ NHẠY ĐỊNH GIÁ']),
      ('VI. DỰ PHÓNG KẾT QUẢ KINH DOANH, ĐỊNH GIÁ & KHUYẾN NGHỊ',['ĐỊNH GIÁ TƯƠNG ĐỐI','ĐỊNH GIÁ NỘI TẠI / CƠ SỞ','BEAR - BASE - BULL','KẾT LUẬN PHÂN TÍCH']),
      ('VII. TRIỂN VỌNG NGÀNH',['BỐI CẢNH VĨ MÔ','TRIỂN VỌNG NGÀNH']),
    ]


def _subsection_content(doc,ticker,sub,meta,s,val,rr,report_type):
    desc=dict(professional_page_plan(report_type)).get(sub,'')
    pars=_page_narrative(ticker,sub,desc,meta,s,val,rr,report_type)
    pars.extend(_decision_narrative(ticker,sub,meta,s,val,report_type))
    _compact_pars(doc,pars,max_pars=4 if sub not in ('RỦI RO VĨ MÔ','RỦI RO NGÀNH') else 3)
    if report_type=='analysis' and sub=='BEAR - BASE - BULL': _add_fv_table(doc,ticker)
    elif report_type=='analysis' and sub=='ĐỊNH GIÁ NỘI TẠI / CƠ SỞ': _add_triangulation_table(doc,ticker)
    elif report_type=='rating' and sub in ('ANCHOR','ĐIỀU CHỈNH NỘI SINH / MODIFIERS','SACP / SCA','ICR'):
        filters={'ANCHOR':['BICRA / Anchor'],'ĐIỀU CHỈNH NỘI SINH / MODIFIERS':['Hồ sơ Kinh doanh','Vốn và Lợi nhuận','Vị thế Rủi ro','Huy động vốn và Thanh khoản'],'SACP / SCA':['SACP','SCA'],'ICR':['Hỗ trợ bên ngoài','ICR']}
        _add_waterfall(doc,ticker,filters.get(sub),compact=True)
    else: _evidence_table(doc,ticker,sub,report_type)

    metrics=_page_peer_metrics(sub,meta.get('EntityType'))
    imgs=[]
    # historical + peer are paired horizontally; this is the main whitespace fix.
    if metrics:
        m=metrics[0]
        try: imgs.append((f"Xu hướng {VI_METRIC.get(m,METH_LABELS.get(m,m))}",chart_metric(ticker,m)))
        except: pass
        try: imgs.append((f"So sánh peer - {VI_METRIC.get(m,METH_LABELS.get(m,m))}",peer_bar_chart(ticker,m,top_n=10)))
        except: pass
    _add_two_charts(doc,imgs,width_mm=76)


def _compact_appendix(doc,ticker,report_type):
    _section_band(doc,'PHỤ LỤC')
    _subhead(doc,'Ma trận chỉ tiêu theo phương pháp')
    z=methodology_kpi_table(ticker,include_missing=False)
    if not z.empty:
        # one compact table rather than many separate tables/section headings
        tb=doc.add_table(rows=1,cols=5);tb.style='Table Grid';tb.alignment=WD_TABLE_ALIGNMENT.CENTER
        hdr=['Nhóm','Chỉ tiêu','Doanh nghiệp','TB ngành','Trung vị']
        for j,x in enumerate(hdr):tb.rows[0].cells[j].text=x;_set_cell_shading(tb.rows[0].cells[j],LIGHT_GREEN)
        _set_repeat_table_header(tb.rows[0])
        for _,r in z.iterrows():
            c=tb.add_row().cells;m=r['Metric'];vals=[r['Nhóm phân tích'],r['Chỉ tiêu'],_fmt_method_value(m,r['Doanh nghiệp']),_fmt_method_value(m,r['TB ngành']),_fmt_method_value(m,r['Trung vị'])]
            for j,x in enumerate(vals):c[j].text=str(x)
        for row in tb.rows:
            trPr=row._tr.get_or_add_trPr();cant=OxmlElement('w:cantSplit');trPr.append(cant)
            for c in row.cells:
                _set_cell_margins(c,top=22,start=35,bottom=22,end=35)
                for p in c.paragraphs:
                    p.paragraph_format.space_after=Pt(0);p.paragraph_format.line_spacing=1.0
                    for r in p.runs:r.font.name='Lato';r.font.size=Pt(7.5)
    if report_type=='rating':
        _subhead(doc,'Waterfall và audit trail')
        _add_waterfall(doc,ticker,None,compact=False)
    _subhead(doc,'Đồ thị so sánh peer chuyên sâu')
    peer_metrics=metric_list(ticker,available_only=True)
    chartable={'TotalAssets','GrossLoans','CustomerDeposits','Equity','Revenue','ROE','ROA','NIM','NPL','CAR','CIR','LDR','CASA','PB','PE','DebtEquity','CurrentRatio','AvailableCapitalRatio','GrossMargin','NetMargin','DebtEBITDA','CFO_Debt','AssetEquity','CreditCostProxy','FundingGapAssets','CashAssets','WorkingCapitalAssets','NetDebtEquity','NetDebtEBITDA','EquityAssetsCorp','AssetTurnover','FOCFMargin','CashDebt','DebtAssets','FOCF_Debt','CFO_Margin','CapexRevenue'}
    charts=[]
    for mm in [m for m in peer_metrics if m in chartable][:12]:
        try:
            b=peer_bar_chart(ticker,mm,VI_METRIC.get(mm,METH_LABELS.get(mm,mm)),top_n=8)
            if b:charts.append((VI_METRIC.get(mm,METH_LABELS.get(mm,mm)),b))
        except:pass
    # Flow chart pairs naturally. Do not impose atlas page breaks: Word uses every remaining vertical space.
    for k in range(0,len(charts),2):
        batch=charts[k:k+2];tbl=doc.add_table(rows=1,cols=2);tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        for j,(label,bio) in enumerate(batch):
            c=tbl.cell(0,j);_set_cell_margins(c,top=15,start=20,bottom=15,end=20)
            p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(0)
            r=p.add_run(label);r.bold=True;r.font.size=Pt(7.5);r.font.name='Lato'
            q=c.add_paragraph();q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_after=Pt(0);q.add_run().add_picture(bio,width=Mm(76))
        if len(batch)==1:
            tbl.cell(0,1).text=''


def generate_docx(ticker,report_type='analysis',rating_result=None,mna=None):
    ticker=str(ticker).upper();meta=get_company(ticker);s=get_snapshot(ticker);val=valuation(ticker,s)
    rr=rating_result or (rate_company(ticker) if report_type=='rating' else {})
    doc=Document();_style_doc(doc, report_type)
    # Match supplied report density: 10.5pt body, wider text area, green section bands.
    sec=doc.sections[0];sec.top_margin=Mm(13);sec.bottom_margin=Mm(13);sec.left_margin=Mm(14);sec.right_margin=Mm(14)
    n=doc.styles['Normal'];n.font.size=Pt(10.2);n.paragraph_format.line_spacing=1.06;n.paragraph_format.space_after=Pt(3)
    _cover_sample(doc,ticker,meta,report_type)
    if report_type=='rating':
        _rating_summary_page(doc,ticker,meta,rr)
        # KLB/VDS samples flow directly from the rating summary into rating drivers; no sparse TOC page.
        groups=_rating_groups(meta.get('EntityType'))
    else:
        _analysis_summary_page(doc,ticker,meta,s,val)
        groups=_analysis_groups(meta.get('EntityType'))
        # ASEANSC-style compact TOC placed in the remaining space of the summary page when possible.
        _section_band(doc,'MỤC LỤC')
        for i,(group,_) in enumerate(groups,1):
            p=doc.add_paragraph();p.paragraph_format.space_after=Pt(1);r=p.add_run(group);r.bold=True;r.font.name='Lato';r.font.size=Pt(9.5)

    for group,subs in groups:
        _section_band(doc,group)
        for sub in subs:
            _subhead(doc,sub)
            _subsection_content(doc,ticker,sub,meta,s,val,rr,report_type)

    _compact_appendix(doc,ticker,report_type)

    # Final pagination discipline: no forced page breaks between normal sections; prevent orphan headings/table rows.
    for p in doc.paragraphs:
        if p.style and str(p.style.name).startswith('Heading'): p.paragraph_format.keep_with_next=True
        p.paragraph_format.widow_control=True
    for table in doc.tables:
        for row in table.rows:
            trPr=row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None: trPr.append(OxmlElement('w:cantSplit'))
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.widow_control=True
                    for run in p.runs:
                        run.font.name='Lato'
                        if run.font.size is None:run.font.size=Pt(9)
    bio=BytesIO();doc.save(bio);return bio.getvalue()

def generate_pdf(ticker,report_type='analysis',rating_result=None,mna=None):
    """PDF is converted from the same DOCX source so DOCX/PDF layout stays identical.
    Falls back to a minimal ReportLab PDF only when LibreOffice is unavailable.
    """
    import tempfile, subprocess, shutil, os
    docx_bytes=generate_docx(ticker,report_type,rating_result,mna)
    soffice=shutil.which('libreoffice') or shutil.which('soffice')
    if soffice:
        with tempfile.TemporaryDirectory() as td:
            inp=Path(td)/'report.docx'; inp.write_bytes(docx_bytes)
            env=os.environ.copy(); env['HOME']=td
            cmd=[soffice,'--headless','--convert-to','pdf','--outdir',td,str(inp)]
            try:
                subprocess.run(cmd,check=True,stdout=subprocess.PIPE,stderr=subprocess.PIPE,timeout=90,env=env)
                pdf=Path(td)/'report.pdf'
                if pdf.exists(): return pdf.read_bytes()
            except Exception: pass
    # Minimal fallback only; normal local/cloud deployment should install LibreOffice via packages.txt.
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    bio=BytesIO();pdf=SimpleDocTemplate(bio,pagesize=A4);styles=getSampleStyleSheet()
    meta=get_company(str(ticker).upper())
    story=[Paragraph('BÁO CÁO XẾP HẠNG TÍN NHIỆM' if report_type=='rating' else 'BÁO CÁO PHÂN TÍCH CỔ PHIẾU',styles['Title']),Spacer(1,12),Paragraph(f"{str(ticker).upper()} - {meta.get('CompanyName')}",styles['Heading1']),Paragraph('PDF fallback: vui lòng cài LibreOffice để PDF giữ nguyên format DOCX.',styles['BodyText'])]
    pdf.build(story);return bio.getvalue()


# ===== V8.40 SAMPLE-NATIVE REPORT ENGINE =====
# Goal: mirror the supplied analyst/rating reports: evidence first, narrative + chart/table integrated,
# methodology moved to the back, and major-section order follows the reference reports.

def _v840_value(metric, v):
    x=num(v)
    if x is None:return 'N/A'
    if metric in PCT or metric in METH_PCT:return pct(x)
    if metric in MULT or metric in METH_MULT:return mult(x)
    if abs(x)>=1e12:return (f"{x/1e12:,.1f} nghìn tỷ").replace(',','X').replace('.',',').replace('X','.')
    if abs(x)>=1e9:return (f"{x/1e9:,.0f} tỷ").replace(',','X').replace('.',',').replace('X','.')
    if abs(x)>=1e6:return (f"{x/1e6:,.1f} triệu").replace(',','X').replace('.',',').replace('X','.')
    return vi(x,2)

def _v840_metric_row(ticker,metric):
    k,_,_=sector_kpi_table(ticker)
    if k is None or not len(k): return None
    r=k[k.Metric.astype(str).eq(metric)]
    if not len(r):return None
    return r.iloc[-1]

def _v840_fact(ticker,metric):
    r=_v840_metric_row(ticker,metric)
    if r is None:return ''
    c=num(r.get('Doanh nghiệp')); m=num(r.get('Trung bình ngành')); med=num(r.get('Trung vị ngành')); n=int(r.get('Số DN có dữ liệu',0) or 0)
    if c is None:return ''
    label=str(r.get('Chỉ tiêu') or VI_METRIC.get(metric,METH_LABELS.get(metric,metric)))
    txt=f"{label} đạt {_v840_value(metric,c)}"
    if m is not None:
        if metric in {'NPL','CIR','DebtEquity','DebtEBITDA','NetDebtEBITDA','CreditCostProxy'}:
            better='thấp hơn' if c<m else 'cao hơn'
        else: better='cao hơn' if c>m else 'thấp hơn'
        txt+=f", {better} mức bình quân peer {_v840_value(metric,m)}"
    if med is not None and m is not None and abs(m-med)/(abs(m)+1e-9)>0.08:
        txt+=f" (trung vị {_v840_value(metric,med)})"
    if n:txt+=f"; mẫu so sánh {n} doanh nghiệp"
    return txt+'.'

def _v840_trend_fact(ticker,metric):
    h=entity_history(ticker)
    if h is None or not len(h):return ''
    z=h[h.Metric.astype(str).eq(metric)].copy()
    if not len(z):return ''
    z['Date']=z.Period.map(_period_date);z['Value']=pd.to_numeric(z.Value,errors='coerce')
    z=z.dropna(subset=['Date','Value']).sort_values('Date').drop_duplicates('Date',keep='last')
    if len(z)<2:return ''
    a=float(z.iloc[-1].Value);b=float(z.iloc[-2].Value)
    label=VI_METRIC.get(metric,METH_LABELS.get(metric,metric));period=str(z.iloc[-1].Period)
    if b==0:return ''
    if metric in PCT or metric in METH_PCT:
        ch=(a-b)*100; direction='tăng' if ch>0 else 'giảm'
        return f"So với kỳ liền trước, {label} {direction} {abs(ch):.1f} điểm % lên {_v840_value(metric,a)} tại {period}.".replace('.',',',1) if False else f"So với kỳ liền trước, {label} {direction} {str(round(abs(ch),1)).replace('.',',')} điểm % lên {_v840_value(metric,a)} tại {period}."
    ch=a/b-1;direction='tăng' if ch>0 else 'giảm'
    return f"So với kỳ liền trước, {label} {direction} {str(round(abs(ch)*100,1)).replace('.',',')}% lên {_v840_value(metric,a)} tại {period}."

def _v840_cross_insights(ticker,metrics):
    vals={}
    for m in metrics:
        r=_v840_metric_row(ticker,m)
        if r is not None: vals[m]=num(r.get('Doanh nghiệp'))
    out=[]
    if all(vals.get(k) is not None for k in ['ROA','AssetEquity','ROE']):
        implied=vals['ROA']*vals['AssetEquity']
        out.append(f"ROE có thể được đọc cùng ROA và đòn bẩy tài sản: ROA {_v840_value('ROA',vals['ROA'])} × Tổng tài sản/VCSH {_v840_value('AssetEquity',vals['AssetEquity'])} hàm ý ROE xấp xỉ {_v840_value('ROE',implied)}, so với ROE ghi nhận {_v840_value('ROE',vals['ROE'])}.")
    if all(vals.get(k) is not None for k in ['NIM','CASA','LDR']):
        out.append(f"NIM {_v840_value('NIM',vals['NIM'])} cần được đọc cùng CASA {_v840_value('CASA',vals['CASA'])} và LDR {_v840_value('LDR',vals['LDR'])}; CASA cao hỗ trợ chi phí vốn, trong khi LDR cao làm giảm dư địa thanh khoản và có thể gây áp lực huy động.")
    if all(vals.get(k) is not None for k in ['NPL','CAR']):
        out.append(f"NPL {_v840_value('NPL',vals['NPL'])} và CAR {_v840_value('CAR',vals['CAR'])} phản ánh đồng thời rủi ro tổn thất kỳ vọng và năng lực hấp thụ lỗ; kết hợp hai chỉ tiêu này đáng tin cậy hơn việc đánh giá riêng từng tỷ lệ.")
    if all(vals.get(k) is not None for k in ['DebtEquity','CurrentRatio']):
        out.append(f"Đòn bẩy {_v840_value('DebtEquity',vals['DebtEquity'])} đi cùng hệ số thanh toán hiện hành {_v840_value('CurrentRatio',vals['CurrentRatio'])}; mức đòn bẩy chỉ bền vững khi thanh khoản và dòng tiền đủ để đáp ứng nghĩa vụ ngắn hạn.")
    return out

def _v840_analysis_text(ticker,metrics,context=''):
    facts=[]
    for m in metrics:
        f=_v840_fact(ticker,m)
        if f:facts.append(f)
    trends=[]
    for m in metrics[:2]:
        f=_v840_trend_fact(ticker,m)
        if f:trends.append(f)
    out=[]
    if facts: out.append(' '.join(facts[:5]))
    if trends: out.append(' '.join(trends[:3]))
    out.extend(_v840_cross_insights(ticker,metrics)[:2])
    if context: out.append(context)
    return out

def _v840_mini_table(cell,ticker,metrics):
    rows=[]
    for m in metrics:
        r=_v840_metric_row(ticker,m)
        if r is not None:rows.append((m,r))
    if not rows:return
    t=cell.add_table(rows=1,cols=4);t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,x in enumerate(['Chỉ tiêu','DN','TB peer','Trung vị']):
        t.cell(0,j).text=x;_set_cell_shading(t.cell(0,j),LIGHT_GREEN)
    for m,r in rows:
        c=t.add_row().cells
        vals=[str(r.get('Chỉ tiêu')), _v840_value(m,r.get('Doanh nghiệp')),_v840_value(m,r.get('Trung bình ngành')),_v840_value(m,r.get('Trung vị ngành'))]
        for j,x in enumerate(vals):c[j].text=str(x)
    for row in t.rows:
        for c in row.cells:
            _set_cell_margins(c,top=20,start=30,bottom=20,end=30)
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(0);p.paragraph_format.line_spacing=1.0
                for r in p.runs:r.font.name='Lato';r.font.size=Pt(7.4)

def _v840_integrated_block(doc,ticker,title,metrics,paras=None,chart=None,reverse=False):
    """Narrative + evidence on the same visual block, as in the user's sample reports."""
    _subhead(doc,title)
    tbl=doc.add_table(rows=1,cols=2);tbl.autofit=False;tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    left,right=tbl.cell(0,0),tbl.cell(0,1)
    left.width=Mm(98);right.width=Mm(72)
    _set_cell_margins(left,top=25,start=20,bottom=20,end=55);_set_cell_margins(right,top=25,start=55,bottom=20,end=20)
    textcell,viscell=(right,left) if reverse else (left,right)
    if paras is None:paras=_v840_analysis_text(ticker,metrics)
    # overwrite default empty para first
    p=textcell.paragraphs[0];p.clear()
    for i,txt in enumerate([x for x in paras if x]):
        p=textcell.paragraphs[0] if i==0 else textcell.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.line_spacing=1.05;p.paragraph_format.space_after=Pt(4)
        r=p.add_run(str(txt));r.font.name='Lato';r.font.size=Pt(9.2)
    # chart first, then compact table immediately below it
    cm=chart or (metrics[0] if metrics else None)
    if cm:
        try:
            bio=chart_metric(ticker,cm,percent=cm in PCT or cm in METH_PCT)
            q=viscell.paragraphs[0];q.alignment=WD_ALIGN_PARAGRAPH.CENTER;q.paragraph_format.space_after=Pt(1)
            q.add_run().add_picture(bio,width=Mm(69))
        except Exception:pass
    _v840_mini_table(viscell,ticker,metrics[:8])
    _set_cell_border(left,bottom={'val':'single','sz':'4','color':'D9D9D9'});_set_cell_border(right,bottom={'val':'single','sz':'4','color':'D9D9D9'})

def _v840_rating_groups(entity_type):
    if entity_type=='BANK':
        return [
          ('NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG','drivers'),
          ('THÔNG TIN TỔNG QUAN TỔ CHỨC PHÁT HÀNH','overview'),
          ('RỦI RO VĨ MÔ','macro'),('RỦI RO NGÀNH','industry'),('HỒ SƠ KINH DOANH','business'),
          ('VỐN VÀ LỢI NHUẬN','capital_profit'),('VỊ THẾ RỦI RO','risk_position'),
          ('HUY ĐỘNG VỐN VÀ THANH KHOẢN','funding_liquidity'),('YẾU TỐ BÊN NGOÀI','support'),
          ('ĐỘ NHẠY XẾP HẠNG','sensitivity')]
    if entity_type=='SECURITIES':
        return [
          ('NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG','drivers'),('THÔNG TIN TỔNG QUAN TỔ CHỨC PHÁT HÀNH','overview'),
          ('RỦI RO VĨ MÔ','macro'),('RỦI RO NGÀNH','industry'),('HỒ SƠ KINH DOANH','business'),
          ('VỐN, ĐÒN BẨY VÀ LỢI NHUẬN','capital_profit'),('VỊ THẾ RỦI RO','risk_position'),
          ('NGUỒN VỐN VÀ THANH KHOẢN','funding_liquidity'),('YẾU TỐ BÊN NGOÀI','support'),('ĐỘ NHẠY XẾP HẠNG','sensitivity')]
    return [
      ('NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG','drivers'),('THÔNG TIN TỔNG QUAN TỔ CHỨC PHÁT HÀNH','overview'),
      ('RỦI RO VĨ MÔ VÀ NGÀNH','macro_industry'),('HỒ SƠ KINH DOANH','business'),('RỦI RO TÀI CHÍNH','financial_risk'),
      ('QUẢN TRỊ VÀ QUẢN LÝ','governance'),('THANH KHOẢN','funding_liquidity'),('YẾU TỐ BÊN NGOÀI','support'),('ĐỘ NHẠY XẾP HẠNG','sensitivity')]

def _v840_driver_page(doc,ticker,meta,rr):
    a=intelligent_analyze(ticker)
    _section_band(doc,'NHỮNG NHÂN TỐ CHÍNH DẪN ĐẾN KẾT QUẢ XẾP HẠNG')
    # Use actual strengths/risks from the engine plus live KPI evidence; no methodology exposition.
    t=doc.add_table(rows=1,cols=2);t.autofit=False
    for idx,(head,items) in enumerate([('ĐIỂM MẠNH',a.get('Strengths',[])[:4]),('ĐIỂM HẠN CHẾ / RỦI RO',a.get('Risks',[])[:4])]):
        c=t.cell(0,idx);_set_cell_margins(c,top=60,start=70,bottom=60,end=70);_set_cell_shading(c,'F5FAF0' if idx==0 else 'FAF7F2')
        p=c.paragraphs[0];r=p.add_run(head);r.bold=True;r.font.name='Lato';r.font.size=Pt(10);r.font.color.rgb=None
        for x in items:
            p=c.add_paragraph(style=None);p.paragraph_format.space_after=Pt(3);p.paragraph_format.left_indent=Mm(2)
            r=p.add_run('• '+str(x));r.font.name='Lato';r.font.size=Pt(9)
    et=meta.get('EntityType')
    metrics={'BANK':['ROE','NPL','CAR','CASA'],'SECURITIES':['ROE','AvailableCapitalRatio','DebtEquity','CurrentRatio']}.get(et,['ROE','DebtEquity','CurrentRatio','CFO_Debt'])
    facts=_v840_analysis_text(ticker,metrics)
    for ptxt in facts:
        p=doc.add_paragraph(ptxt);p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.space_after=Pt(3)
    _v840_mini_table(doc.add_table(rows=1,cols=1).cell(0,0),ticker,metrics)

def _v840_rating_body(doc,ticker,meta,s,rr):
    et=meta.get('EntityType')
    for heading,key in _v840_rating_groups(et):
        if key=='drivers':
            _v840_driver_page(doc,ticker,meta,rr);continue
        _section_band(doc,heading)
        if key=='overview':
            pars=[f"{company_display_name(meta,ticker)} hoạt động trong ngành {meta.get('Sector')}. Quy mô và vị trí tương đối được đánh giá trực tiếp qua tổng tài sản/doanh thu, vốn chủ sở hữu và các chỉ tiêu hoạt động chính so với peer."]
            ms=['TotalAssets','GrossLoans','CustomerDeposits','LoanAssets','DepositAssets','AssetEquity'] if et=='BANK' else (['TotalAssets','Revenue','Equity'] if et=='SECURITIES' else ['Revenue','TotalAssets','Equity'])
            _v840_integrated_block(doc,ticker,'Quy mô hoạt động và vị trí tương đối',ms,pars+_v840_analysis_text(ticker,ms),reverse=False);continue
        if key in ('macro','industry','macro_industry'):
            only='macro' if key=='macro' else 'industry' if key=='industry' else None
            intel=_public_intel_paragraphs(et,only=only)
            for title,nar,source,url,asof in intel[:4]:
                if title:_subhead(doc,title)
                _intel_body(doc,nar)
                _intel_source(doc,source,asof,url)
            # Keep macro/industry sections focused on the external environment.
            # Entity-specific transmission analysis belongs in HỒ SƠ KINH DOANH,
            # where it can be read together with the company's actual KPIs.
            continue
        if key=='business':
            ms=['TotalAssets','GrossLoans','CustomerDeposits','LoanAssets','DepositAssets','AssetEquity'] if et=='BANK' else (['Revenue','TotalAssets','ROE'] if et=='SECURITIES' else ['Revenue','GrossMargin','AssetTurnover'])
            ctx="Quy mô chỉ tạo lợi thế khi đi cùng tăng trưởng có chất lượng và khả năng duy trì thị phần. Vì vậy, đánh giá Hồ sơ Kinh doanh ưu tiên khoảng cách với peer và xu hướng nhiều kỳ thay vì chỉ nhìn quy mô tuyệt đối."
            _v840_integrated_block(doc,ticker,'Quy mô, tăng trưởng và vị thế cạnh tranh',ms,_v840_analysis_text(ticker,ms,ctx),chart=ms[0])

            # Move the former 'Liên hệ với hồ sơ doanh nghiệp' block here.
            # This avoids interrupting RỦI RO VĨ MÔ / RỦI RO NGÀNH with issuer-specific KPIs.
            if et=='BANK':
                link_ms=['NIM','CASA','LDR','CAR','NPL']
                link_ctx='Trong Hồ sơ Kinh doanh, tác động của môi trường ngành được đối chiếu trực tiếp với khả năng tạo biên lãi, chất lượng nguồn vốn, mức sử dụng vốn, bộ đệm vốn và chất lượng tài sản của ngân hàng. NIM, CASA, LDR, CAR và NPL được đọc đồng thời với peer để đánh giá mức độ chuyển hóa vị thế kinh doanh thành hiệu quả và sức chống chịu.'
            elif et=='SECURITIES':
                link_ms=['AvailableCapitalRatio','DebtEquity','CurrentRatio','ROE','MarketShareBrokerage']
                link_ctx='Trong Hồ sơ Kinh doanh, tác động của thị trường được đối chiếu với thị phần môi giới, năng lực vốn, đòn bẩy, thanh khoản và khả năng sinh lời để đánh giá khả năng chuyển hóa cơ hội thị trường thành tăng trưởng bền vững.'
            else:
                link_ms=['Revenue','GrossMargin','DebtEquity','CurrentRatio','CFO_Debt']
                link_ctx='Trong Hồ sơ Kinh doanh, tác động của môi trường ngành được đối chiếu với tăng trưởng doanh thu, biên lợi nhuận, cơ cấu vốn, thanh khoản và khả năng tạo dòng tiền để đánh giá sức cạnh tranh và khả năng thích ứng của doanh nghiệp.'
            # Issuer-specific KPI transmission block belongs ONLY in HỒ SƠ KINH DOANH.
            # Do not render this block in RỦI RO VĨ MÔ / RỦI RO NGÀNH to avoid duplicated analysis.
            _v840_integrated_block(doc,ticker,'Liên hệ với hồ sơ doanh nghiệp',link_ms,_v840_analysis_text(ticker,link_ms,link_ctx),chart=link_ms[0])
            continue
        if key=='capital_profit':
            ms=['CAR','EquityAssets','TangibleEquityAssets','AssetEquity','ROE','ROA','ProfitAssets','NIM'] if et=='BANK' else (['ROE','ROA','AvailableCapitalRatio','DebtEquity'] if et=='SECURITIES' else ['ROE','DebtEquity','DebtEBITDA','CFO_Debt'])
            ctx="Khả năng tạo lợi nhuận và bộ đệm vốn được đọc đồng thời: lợi nhuận cao hỗ trợ tích lũy vốn, nhưng tốc độ tăng tài sản/nợ nhanh có thể làm suy giảm vùng đệm nếu vốn nội sinh không theo kịp."
            _v840_integrated_block(doc,ticker,'Khả năng sinh lời và bộ đệm vốn',ms,_v840_analysis_text(ticker,ms,ctx),chart=ms[0]);continue
        if key=='risk_position':
            ms=['NPL','CreditCostProxy','ProvisionOperatingIncome','LoanAssets','CAR','EquityAssets'] if et=='BANK' else (['MarginLoansEquity','DebtEquity','ROA','CurrentRatio'] if et=='SECURITIES' else ['DebtEBITDA','NetDebtEBITDA','CFO_Debt','FOCF_Debt'])
            ctx="Đây là nhóm chỉ tiêu có khả năng truyền dẫn trực tiếp sang lợi nhuận, vốn và khả năng thực hiện nghĩa vụ nợ; chênh lệch bất lợi so với peer được coi là tín hiệu cần giám sát chặt hơn."
            _v840_integrated_block(doc,ticker,'Chất lượng tài sản / khẩu vị rủi ro',ms,_v840_analysis_text(ticker,ms,ctx),chart=ms[0]);continue
        if key in ('funding_liquidity','financial_risk'):
            ms=['CASA','LDR','CustomerDeposits','DepositAssets','FundingGapAssets','CAR'] if et=='BANK' else (['CurrentRatio','DebtEquity','CFO_Debt','CashAssets'] if et=='SECURITIES' else ['CurrentRatio','CFO_Debt','FOCF_Debt','CashDebt'])
            ctx="Thanh khoản được đánh giá trên cả cấu trúc nguồn vốn và khả năng tạo tiền. Một tỷ lệ thanh khoản tốt tại một thời điểm không đủ bù cho cấu trúc đáo hạn tập trung hoặc dòng tiền hoạt động yếu."
            _v840_integrated_block(doc,ticker,'Nguồn vốn, thanh khoản và khả năng trả nợ',ms,_v840_analysis_text(ticker,ms,ctx),chart=ms[0]);continue
        if key=='governance':
            p=doc.add_paragraph("Phần quản trị chỉ trình bày các phát hiện định tính có bằng chứng trong hồ sơ doanh nghiệp: cấu trúc quản trị, chính sách tài chính, quản trị rủi ro, giao dịch bên liên quan và mức độ phụ thuộc hệ sinh thái. Các nội dung chưa có dữ liệu cấu trúc được giữ là điểm cần chuyên viên cập nhật, không tự suy diễn.");p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;continue
        if key=='support':
            p=doc.add_paragraph(f"Kết quả mô hình ghi nhận mức hỗ trợ bên ngoài {rr.get('ExternalSupportNotches',0)} bậc. Tác động hỗ trợ chỉ được đưa vào xếp hạng cuối cùng khi có bằng chứng về năng lực và động cơ hỗ trợ; chi tiết tính toán được đưa xuống phụ lục.");p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;continue
        if key=='sensitivity':
            a=intelligent_analyze(ticker);risks=a.get('Risks',[])[:3]
            p=doc.add_paragraph(f"Bậc xếp hạng hiện tại theo mô hình là {rr.get('ICR','N/A')}. Các yếu tố có thể tạo áp lực hạ bậc gồm: "+('; '.join(risks) if risks else 'suy giảm đáng kể về vốn, chất lượng tài sản, thanh khoản hoặc khả năng sinh lời so với peer.'))
            p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


def _v840_analysis_body(doc,ticker,meta,s,val):
    et=meta.get('EntityType')
    groups=[
      ('I. TỔNG QUAN','overview'),('II. HOẠT ĐỘNG KINH DOANH','operations'),('III. TÌNH HÌNH TÀI CHÍNH','financials'),
      ('IV. LUẬN ĐIỂM ĐẦU TƯ','thesis'),('V. RỦI RO','risks'),('VI. DỰ PHÓNG KẾT QUẢ KINH DOANH, ĐỊNH GIÁ & KHUYẾN NGHỊ','valuation'),
      ('VII. TRIỂN VỌNG NGÀNH','industry')]
    a=intelligent_analyze(ticker)
    for heading,key in groups:
        _section_band(doc,heading)
        if key=='overview':
            ms=['TotalAssets','GrossLoans','CustomerDeposits'] if et=='BANK' else ['Revenue','TotalAssets','Equity']
            _v840_integrated_block(doc,ticker,'Quy mô và vị thế doanh nghiệp',ms,_v840_analysis_text(ticker,ms,"Quy mô được đặt cạnh tăng trưởng và vị trí peer để xác định lợi thế franchise có thực sự chuyển hóa thành hiệu quả kinh doanh hay không."),chart=ms[0]);continue
        if key=='operations':
            if et=='BANK':
                blocks=[('Thu nhập và khả năng sinh lời',['ROE','ROA','NIM','CIR','NII_OperatingIncome','ProfitAssets'],'ROE'),('Tăng trưởng tín dụng và nguồn vốn',['GrossLoans','CustomerDeposits','LDR','CASA','LoanAssets','DepositAssets'],'GrossLoans')]
            elif et=='SECURITIES':blocks=[('Tăng trưởng và hiệu quả hoạt động',['Revenue','ROE','ROA','NetMargin'],'Revenue'),('Cơ cấu vốn phục vụ kinh doanh',['AvailableCapitalRatio','DebtEquity','CurrentRatio'],'AvailableCapitalRatio')]
            else:blocks=[('Tăng trưởng và biên lợi nhuận',['Revenue','GrossMargin','NetMargin','AssetTurnover'],'Revenue'),('Hiệu quả vốn',['ROE','ROA','DebtEquity'],'ROE')]
            for i,(t,ms,ch) in enumerate(blocks):
                _v840_integrated_block(doc,ticker,t,ms,_v840_analysis_text(ticker,ms),chart=ch,reverse=bool(i%2))
            # Issuer-specific KPI transmission belongs in the business/operations section,
            # not in the industry outlook. Render it once here only.
            if et=='BANK':
                _v840_integrated_block(
                    doc,ticker,'Hiệu quả kinh doanh và sức chống chịu',
                    ['NIM','CASA','LDR','NPL','CAR','ROE'],
                    _v840_analysis_text(
                        ticker,['NIM','CASA','LDR','NPL','CAR','ROE'],
                        'Các chỉ tiêu trên cho thấy mức độ doanh nghiệp có thể hưởng lợi từ tăng trưởng tín dụng mà không đánh đổi quá mức NIM, chất lượng tài sản hoặc thanh khoản.'
                    ),
                    chart='NIM'
                )
            continue
        if key=='financials':
            if et=='BANK':blocks=[('Chất lượng tài sản và chi phí rủi ro',['NPL','CreditCostProxy','ProvisionOperatingIncome','LoanAssets'],'NPL'),('An toàn vốn và cấu trúc bảng cân đối',['CAR','EquityAssets','TangibleEquityAssets','AssetEquity','ProfitAssets'],'CAR'),('Nguồn vốn và thanh khoản',['CASA','LDR','DepositAssets','FundingGapAssets','CustomerDeposits'],'CASA')]
            elif et=='SECURITIES':blocks=[('Đòn bẩy và thanh khoản',['DebtEquity','CurrentRatio','CFO_Debt','CashAssets'],'DebtEquity'),('Rủi ro bảng cân đối',['MarginLoansEquity','AvailableCapitalRatio','ROA'],'MarginLoansEquity')]
            else:blocks=[('Đòn bẩy và khả năng trả nợ',['DebtEquity','DebtEBITDA','CFO_Debt','FOCF_Debt'],'DebtEquity'),('Thanh khoản và dòng tiền',['CurrentRatio','CashDebt','FOCFMargin'],'CurrentRatio')]
            for i,(t,ms,ch) in enumerate(blocks):_v840_integrated_block(doc,ticker,t,ms,_v840_analysis_text(ticker,ms),chart=ch,reverse=bool(i%2));continue
        if key=='thesis':
            texts=[]
            if a.get('Strengths'):texts += ['• '+x for x in a['Strengths'][:4]]
            if a.get('Conclusion'):texts.insert(0,a['Conclusion'])
            for x in texts:
                p=doc.add_paragraph(x);p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY;p.paragraph_format.space_after=Pt(3)
            try:
                sc=peer_scatter_chart(ticker,'ROE','PB')
                if sc:doc.add_picture(sc,width=Mm(112))
            except:pass
            continue
        if key=='risks':
            risks=a.get('Risks',[])[:5]
            for x in risks:
                p=doc.add_paragraph('• '+str(x));p.paragraph_format.space_after=Pt(3)
            if et=='BANK':_v840_integrated_block(doc,ticker,'Áp lực NIM, chất lượng tài sản và thanh khoản',['NIM','NPL','LDR','CAR'],_v840_analysis_text(ticker,['NIM','NPL','LDR','CAR']),chart='NIM')
            continue
        if key=='valuation':
            _v840_integrated_block(doc,ticker,'So sánh định giá với peer',['PB','PE','ROE','ROA'],_v840_analysis_text(ticker,['PB','PE','ROE','ROA'],"Premium/discount định giá chỉ có ý nghĩa khi được đặt cạnh ROE, tăng trưởng và rủi ro tương đối."),chart='PB')
            _subhead(doc,'Dự phóng và vùng giá')
            _add_fv_table(doc,ticker)
            vt=triangulate(ticker);p=doc.add_paragraph(f"Độ tin cậy phân tích hiện tại: {vt.get('AnalyticalConfidence','N/A')}. Vùng giá Bear-Base-Bull phản ánh độ nhạy của giả định, còn giá trị Strategic/M&A được trình bày riêng để tránh cộng premium hai lần.");p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            continue
        if key=='industry':
            intel=_public_intel_paragraphs(et,only='industry') + _public_intel_paragraphs(et,only='macro')
            for title,nar,source,url,asof in intel[:4]:
                if title:_subhead(doc,title)
                _intel_body(doc,nar)
                _intel_source(doc,source,asof,url)

def _v840_appendix(doc,ticker,report_type):
    """No repeated KPI appendix in the client-facing report.

    KPI/peer evidence is embedded in the relevant analytical sections. Detailed
    scorecards, waterfall and audit trail remain available in platform data rather
    than consuming report pages or repeating evidence out of context.
    """
    return


def _normalize_report_typography(doc):
    """User typography standard: Lato 11pt body; 6pt before/0pt after; tables/charts 10pt."""
    # Normal/body style. Keep designed cover/section headings larger, but all ordinary text is 11pt.
    normal = doc.styles['Normal']
    normal.font.name = 'Lato'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(6)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # Body paragraphs outside tables. Preserve intentionally larger title/heading typography.
    table_paragraph_ids = {id(pp) for tb in doc.tables for row in tb.rows for cell in row.cells for pp in cell.paragraphs}
    for p in doc.paragraphs:
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        style_name = (p.style.name if p.style else '') or ''
        preserve_large = style_name in ('Title', 'Subtitle') or style_name.startswith('Heading')
        for r in p.runs:
            r.font.name = 'Lato'
            if not preserve_large:
                r.font.size = Pt(11)

    # Every table, including summary/KPI/appendix tables: Lato 10pt. Compact cell spacing.
    for tb_idx, tb in enumerate(doc.tables):
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    for r in p.runs:
                        r.font.name = 'Lato'
                        # The first table is the cover's green title band. Preserve its
                        # deliberately larger 30pt/22pt hierarchy; analytical tables remain 10pt.
                        if tb_idx != 0:
                            r.font.size = Pt(10)

    # Matplotlib charts are generated at 10pt globally; reinforce here for subsequent charts.
    plt.rcParams.update({'font.family': 'Lato', 'font.size': 10})

def generate_docx(ticker,report_type='analysis',rating_result=None,mna=None):
    ticker=str(ticker).upper();meta=get_company(ticker);s=get_snapshot(ticker);val=valuation(ticker,s)
    rr=rating_result or (rate_company(ticker) if report_type=='rating' else {})
    doc=Document();_style_doc(doc, report_type)
    sec=doc.sections[0];sec.top_margin=Mm(12);sec.bottom_margin=Mm(12);sec.left_margin=Mm(13);sec.right_margin=Mm(13)
    n=doc.styles['Normal'];n.font.size=Pt(9.6);n.paragraph_format.line_spacing=1.04;n.paragraph_format.space_after=Pt(2.5)
    _cover_sample(doc,ticker,meta,report_type)
    if report_type=='rating':
        _rating_summary_page(doc,ticker,meta,rr)
        _v840_rating_body(doc,ticker,meta,s,rr)
    else:
        _analysis_summary_page(doc,ticker,meta,s,val)
        _section_band(doc,'MỤC LỤC')
        for h in ['I. TỔNG QUAN','II. HOẠT ĐỘNG KINH DOANH','III. TÌNH HÌNH TÀI CHÍNH','IV. LUẬN ĐIỂM ĐẦU TƯ','V. RỦI RO','VI. DỰ PHÓNG KẾT QUẢ KINH DOANH, ĐỊNH GIÁ & KHUYẾN NGHỊ','VII. TRIỂN VỌNG NGÀNH']:
            p=doc.add_paragraph(h);p.paragraph_format.space_after=Pt(1);p.runs[0].font.size=Pt(9);p.runs[0].bold=True
        _v840_analysis_body(doc,ticker,meta,s,val)
    _v840_appendix(doc,ticker,report_type)
    for p in doc.paragraphs:p.paragraph_format.widow_control=True
    for table in doc.tables:
        for row in table.rows:
            trPr=row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:trPr.append(OxmlElement('w:cantSplit'))
    _normalize_report_typography(doc)
    bio=BytesIO();doc.save(bio);return bio.getvalue()
